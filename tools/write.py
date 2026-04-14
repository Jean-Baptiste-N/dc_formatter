#!/usr/bin/env python3

from argparse import ArgumentParser

import docx
from docx.document import Document
from docx.enum.table import WD_ALIGN_VERTICAL
# from docx.enum.text import WD_COLOR
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from tools.utils import get_font_props
from tools.utils import get_format_props
from tools.utils import get_ilvl
from tools.utils import rec_add_xml_children
from tools.utils import TEMPLATE_DC_PATH


def write_simplified_runs(src_p: Paragraph, dst_p: Paragraph) -> None:
    """
    Generate a simplified version of paragraph `src_p`'s list of runs by merging successive runs
    with the same font properties and squashing those containing a single whitespace character.
    Write the resulting list of runs into paragraph `dst_p`.
    """
    # Exclude runs with no text
    runs = [r for r in src_p.runs if r.text]
    # Remove any run sharing all its font properties with the previous run
    i = 1
    while i < len(runs):
        run_font = get_font_props(runs[i], src_p)
        previous_run_font = get_font_props(runs[i - 1], src_p)
        if run_font == previous_run_font or runs[i].text == " ":
            runs[i - 1].text += runs.pop(i).text
        else:
            i += 1

    for r in runs:
        run_font = get_font_props(r, src_p)
        dst_run = dst_p.add_run(text=r.text)
        for prop in ("name", "size"):
            if run_font[prop] is not None:
                setattr(dst_run.font, prop, run_font[prop])
        if run_font["color"] is not None:
            dst_run.font.color.rgb = RGBColor.from_string(run_font["color"])
        # dst_run.font.highlight_color = WD_COLOR.YELLOW
        for prop in ("bold", "italic", "underline"):
            if run_font[prop] is True:
                setattr(dst_run, prop, run_font[prop])


def write_simplified_table(src_table: Table, dst_dc: Document) -> None:
    """
    Generate a simplified version of `src_table` by rmeoving empty paragraphs and calling
    `write_simplified_runs` on the remaining paragraphs of each cell.
    Write the result as a newly created table into the `dst_dc` document.
    """
    dst_table = dst_dc.add_table(len(src_table.rows), len(src_table.columns))
    # dst_table.autofit = True

    # set table left and right margins to 0, and a negative left indent
    cell_mar_attrs = {"w": "0", "type": "dxa"}
    tbl_cell_mar = {"tblCellMar": {"left": cell_mar_attrs, "right": cell_mar_attrs}}
    rec_add_xml_children(dst_table._tblPr, tbl_cell_mar)
    # rec_add_xml_children(dst_table._tblPr, {"shd": {"fill": "FFFF00", "val": "clear"}})
    # tbl_ind = {"tblInd": {"w": str(Inches(-.2).twips), "type": "dxa"}}
    # rec_add_xml_children(dst_table._tblPr, tbl_ind)
    # dst_table.columns[0].width = Inches(1)
    # dst_table.columns[1].width = Inches(10)
    for i_row, src_row in enumerate(src_table.rows):
        dst_row = dst_table.rows[i_row]
        i_cell = 0
        while i_cell < len(src_row.cells):
            src_cell, dst_cell = src_row.cells[i_cell], dst_row.cells[i_cell]
            dst_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for src_p in (p for p in src_cell.paragraphs if p.text):
                # an empty paragraph is automatically added to every new cell, so we need to use it
                dst_p = (dst_cell.paragraphs[0] if not dst_cell.paragraphs[0].text
                         else dst_cell.add_paragraph())
                fmt = get_format_props(src_p)
                if fmt["alignment"] is not None:
                    dst_p.alignment = fmt["alignment"]

                dst_p.paragraph_format.space_after = 0

                write_simplified_runs(src_p, dst_p)

            # look for merged cells in the current row (cells over multiple rows are not supported)
            if (span := src_cell.grid_span) is not None and span > 1:
                i_cell += span - 1
                dst_cell.merge(dst_row.cells[i_cell])

                # add a solid bottom border to the table
                border_attrs = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
                tbl_borders = {"tblBorders": {"bottom": border_attrs}}
                rec_add_xml_children(dst_table._tblPr, tbl_borders)
                # highlight in yellow
                rec_add_xml_children(dst_cell._tc.tcPr, {"shd": {"fill": "FFFF00"}})

            i_cell += 1


def write_simplified_copy(src_dc: Document, dst_path: str) -> None:
    """
    Generate a simplified copy of `src_dc` and save it to `dst_path`.
    """
    dst_dc = docx.Document(TEMPLATE_DC_PATH)

    # ilvl of the first list paragraph in the document
    first_ilvl = None

    for src_e in src_dc.iter_inner_content():
        if isinstance(src_e, Table):
            write_simplified_table(src_e, dst_dc)

        else:  # Paragraph
            if not src_e.text:
                continue
            dst_p = dst_dc.add_paragraph()
            fmt = get_format_props(src_e)
            if fmt["alignment"] is not None:
                dst_p.alignment = fmt["alignment"]

            # check if the paragraph is part of a list
            if (ilvl := get_ilvl(src_e)) is not None:
                dst_p.paragraph_format.space_after = 0
                first_ilvl = first_ilvl if first_ilvl is not None else ilvl
                numPr = dst_p._p.get_or_add_pPr().get_or_add_numPr()
                numPr.get_or_add_ilvl().set(qn("w:val"), str(ilvl - first_ilvl))
                numPr.get_or_add_numId().set(qn("w:val"), "1")

            write_simplified_runs(src_e, dst_p)

    s = dst_dc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)

    dst_dc.save(dst_path)


def main(src_path: str = "test/default.docx", dst_path: str | None = None) -> None:
    src_dc = docx.Document(src_path)
    write_simplified_copy(src_dc, dst_path or src_path.replace(".docx", "_simplified.docx"))


if __name__ == "__main__":
    parser = ArgumentParser(description="Write a copy of the input .docx file, "
                                        "with simplified and homogeneous format.")

    parser.add_argument("input_file", metavar="input-file", nargs="?", default="test/default.docx",
                        help="Path to the input .docx file.")

    parser.add_argument("-o", "--output-path", metavar="PATH",
                        help="Specify the path to write the output file to.")

    parser_args = parser.parse_args()
    parser_kwargs = {"dst_path": parser_args.output_path}

    main(parser_args.input_file, **parser_kwargs)
