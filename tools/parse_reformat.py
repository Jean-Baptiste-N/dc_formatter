#!/usr/bin/env python3

from argparse import ArgumentParser
from itertools import tee
from typing import Iterator

import docx
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.table import _Cell
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from tools.utils import get_ilvl
from tools.utils import rec_add_xml_children
from tools.utils import TEMPLATE_DC_PATH


BLACK = "000000"
BLUE = "548DD4"
ORANGE = "EC7C30"
GRAY = "808080"

TEXT_WIDTH = Inches(6.27)

DEFAULT_STYLE = {"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "keep_together": False,
                                      "keep_with_next": False, "page_break_before": False,
                                      "left_indent": 0, "space_before": 0, "space_after": 0},
                 "font": {"name": "Arial", "size": Pt(11), "bold": False, "italic": False,
                          "underline": False, "color": BLACK, "small_caps": False}}

SECTION_HEADER_STYLE = {"paragraph_format": {"left_indent": Inches(-.2)},
                        "font": {"size": Pt(16), "bold": True, "underline": True, "color": BLUE}}

PAGE_HEADER_STYLE = {"paragraph_format": {"left_indent": Inches(-.2), "page_break_before": True},
                     "font": SECTION_HEADER_STYLE["font"]}

COMPANY_FONT_NAME = "Georgia"

CANDIDATES_MAIN_SKILLS = ["compétence", "competence"]
CANDIDATES_EDUCATION = ["formation", "diplôme", "diplome", "certification", "langue"]
CANDIDATES_PROFESSIONNAL_EXPERIENCES = ["expérience", "experience"]


DocElm = Paragraph | Table
PairOfBlocks = tuple[list[DocElm], list[DocElm]]


# MARK: -------- UTIL FUNCTIONS


def peek(tee_iterator: Iterator[DocElm], n: int = 1, raise_if_empty: bool = False) -> DocElm | None:
    """
    Return `tee_iterator`'s n-th next value without moving the input forward.

    Return `None` if `tee_iterator` is exhausted and `raise_if_empty` is `False`.
    """
    # Based on itertools.tee documentation's lookahead example
    # Using tee()'s default argument value n=2, see https://github.com/python/cpython/issues/126701
    _, forked_iterator = tee(tee_iterator)
    for _ in range(n):
        ret = next(forked_iterator) if raise_if_empty else next(forked_iterator, None)
    return ret


def skip_empty_paragraphs(it_dc: Iterator[DocElm], raise_if_empty: bool = False) -> DocElm | None:
    """
    Iterate over a document element iterator `it_dc` until an element is found that is not an empty
    paragraph. Return this element without consuming it, or None if a StopIteration error occurs and
    `raise_if_empty` is False.
    """
    e = peek(it_dc, raise_if_empty=raise_if_empty)
    n = 0
    while isinstance(e, Paragraph) and not e.text:
        next(it_dc)  # skip the previously peeked empty paragraph
        e = peek(it_dc, raise_if_empty=raise_if_empty)
        n += 1
    print(f"\033[33m    Skipped {n} empty paragraphs.\033[0m")
    return e


# MARK: -------- CRITERION VALIDATORS


def cv_paragraphs_until_empty(it_dc: Iterator[DocElm]) -> PairOfBlocks:
    """
    Criterion validator that returns as expected block the next succession of non-empty paragraphs.

    If a table is found before any non-empty paragraph, it is returned as unexpected block.
    """
    expected, unexpected = [], []
    e = skip_empty_paragraphs(it_dc, raise_if_empty=True)

    if isinstance(e, Table):
        unexpected.append(next(it_dc))
    while isinstance(e, Paragraph) and e.text:
        expected.append(next(it_dc))
        e = peek(it_dc)

    return expected, unexpected


def cv_text_match(it_dc: Iterator[DocElm], *, candidates: list[str], optional: bool = False
                  ) -> PairOfBlocks:
    """
    Criterion validator that returns as expected block the next non-empty paragraph if its text
    contains one of the `candidates` (case insensitive).

    Otherwise the element is returned as unexpected block.
    """
    expected, unexpected = [], []
    e = skip_empty_paragraphs(it_dc, raise_if_empty=True)

    if isinstance(e, Paragraph) and any(c.lower() in e.text.lower() for c in candidates):
        expected.append(next(it_dc))
    elif not optional:
        unexpected.append(next(it_dc))

    return expected, unexpected


def cv_table(it_dc: Iterator[DocElm], *, optional: bool = False) -> PairOfBlocks:
    """
    Criterion validator that returns as expected block the next element if it is a table.

    Otherwise the element is returned as unexpected block unless `optional` is True
    in which case 2 empty blocks are returned.
    """
    expected, unexpected = [], []
    e = skip_empty_paragraphs(it_dc, raise_if_empty=True)

    if isinstance(e, Table):
        expected.append(next(it_dc))
    elif not optional:
        unexpected.append(next(it_dc))

    return expected, unexpected


def cv_successive_list_elements(it_dc: Iterator[DocElm], *, optional: bool = False) -> PairOfBlocks:
    """
    Criterion validator that returns as expected block the next succession of paragraphs that are
    part of a list, which *roughly* corresponds to paragraphs with ilvl values.
    However, if the first non-empty paragraph found has no ilvl but is directly followed by a
    paragraph with an ilvl, it is still considered as the first element of the list,
    and in this case any other paragraph with no ilvl is also considered part of the list
    if directly followed by a paragraph with an ilvl.
    Each element may be preceded by a single empty paragraph, which is ignored and removed.

    If the first element that is not an empty paragraph, is not part of a list, then it is returned
    as unexpected block unless `optional` is True in which case two empty blocks are returned.
    """
    expected, unexpected = [], []
    e = skip_empty_paragraphs(it_dc, raise_if_empty=True)

    first_ilvl = get_ilvl(e)
    e2 = peek(it_dc, 2)
    e2 = e2 if e2 is None or isinstance(e2, Table) or e2.text else peek(it_dc, 3)
    if (isinstance(e, Table)
        or first_ilvl is None and (isinstance(e2, Table) or get_ilvl(e2) is None)):
        unexpected = [] if optional else [next(it_dc)]
        return expected, unexpected

    while isinstance(e, Paragraph):
        if not e.text:
            next(it_dc)  # allow one empty paragraph
            e = peek(it_dc)
            if not (isinstance(e, Paragraph) and e.text):
                break  # two successive empty paragraphs end the list
        if get_ilvl(e) is None and (first_ilvl is not None or get_ilvl(peek(it_dc, 2)) is None):
            break
        expected.append(next(it_dc))
        e = peek(it_dc)

    return expected, unexpected


def cv_paragraphs_before_list(it_dc: Iterator[DocElm]) -> PairOfBlocks:
    """
    Criterion validator that returns as expected block the next succession of paragraphs that
    directly precedes the start of a list
    (c.f. `cv_successive_list_elements` for the definition of a list).

    If a table is found, it is returned as unexpected block and the paragraphs read so far are
    returned as expected block.

    If the first paragraph read is already part of a list, two empty blocks are returned.
    """
    expected, unexpected = [], []
    e = skip_empty_paragraphs(it_dc, raise_if_empty=True)

    e2 = peek(it_dc, 2)
    e2 = e2 if e2 is None or isinstance(e2, Table) or e2.text else peek(it_dc, 3)
    while get_ilvl(e) is None and get_ilvl(e2) is None:
        if isinstance(e, Table):
            unexpected.append(next(it_dc))
            break
        expected.append(next(it_dc))
        e = skip_empty_paragraphs(it_dc)
        e2 = peek(it_dc, 2)
        e2 = e2 if e2 is None or isinstance(e2, Table) or e2.text else peek(it_dc, 3)

    return expected, unexpected


# MARK: -------- BLOCK DEFINITIONS


BLOCK_DEFINITIONS = {
    # First page
    "h_dc": {
        "criterion": (cv_paragraphs_until_empty, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.CENTER},
                    "font": {"size": Pt(20), "bold": True, "color": BLUE},
                    "empty_paragraphs_after": 2}]
    },

    "trigram": {
        "criterion": (cv_paragraphs_until_empty, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.CENTER},
                    "font": {"size": Pt(20), "bold": True, "color": ORANGE},
                    "empty_paragraphs_after": 2}]
    },

    "role": {
        "criterion": (cv_paragraphs_until_empty, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.CENTER},
                    "font": {"size": Pt(20), "bold": True},
                    "empty_paragraphs_after": 2}]
    },

    "years_of_experience": {
        "criterion": (cv_paragraphs_until_empty, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.CENTER},
                    "font": {"size": Pt(20), "bold": True},
                    "empty_paragraphs_after": 2}]
    },

    # Skills page
    "h_main_skills": {
        "criterion": (cv_text_match, {"candidates": CANDIDATES_MAIN_SKILLS}),
        "styles": [PAGE_HEADER_STYLE]
    },

    "lst_main_skills": {
        "criterion": (cv_successive_list_elements, {"optional": True}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                                         "space_before": Inches(.1), "space_after": 0},
                    "font": {"bold": True, "color": ORANGE},
                    "is_list": True, "empty_paragraphs_after": 1},
                   {"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                                         "space_before": 0, "space_after": 0},
                    "preserve_src_bold": True}]
    },

    "tbl_main_skills": {
        "criterion": (cv_table, {"optional": True}),
        "styles": [{"font": {"bold": True, "color": ORANGE},
                    "empty_paragraphs_after": 1},
                   {"font": {"bold": True}}]
    },

    "h_education_1": {
        "criterion": (cv_text_match, {"candidates": CANDIDATES_EDUCATION}),
        "styles": [SECTION_HEADER_STYLE]
    },

    "tbl_education_1": {
        "criterion": (cv_table, {}),
        "styles": [{"col_widths": {0: Inches(1)}, "empty_paragraphs_after": 1}]
    },

    "h_education_2": {
        "criterion": (cv_text_match, {"candidates": CANDIDATES_EDUCATION, "optional": True}),
        "styles": [SECTION_HEADER_STYLE]
    },

    "tbl_education_2": {
        "criterion": (cv_table, {"optional": True}),
        "styles": [{"col_widths": {0: Inches(1)}, "empty_paragraphs_after": 1}]
    },

    "h_education_3": {
        "criterion": (cv_text_match, {"candidates": CANDIDATES_EDUCATION, "optional": True}),
        "styles": [SECTION_HEADER_STYLE]
    },

    "tbl_education_3": {
        "criterion": (cv_table, {"optional": True}),
        "styles": [{"col_widths": {0: Inches(1)}, "empty_paragraphs_after": 1}]
    },

    # Experiences pages
    "h_experiences": {
        "criterion": (cv_text_match, {"candidates": CANDIDATES_PROFESSIONNAL_EXPERIENCES}),
        "styles": [PAGE_HEADER_STYLE]
    },

    "tbl_company_header": {
        "criterion": (cv_table, {}),
        "styles": [{"paragraph_format": {"space_before": 0, "space_after": 0},
                    "font": {"name": COMPANY_FONT_NAME, "size": Pt(20), "color": GRAY},
                    "col_widths": {1: Inches(1.5)}, "side_margins": 0, "bottom_border": True},
                   {"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.RIGHT,
                                         "space_before": 0, "space_after": 0},
                    "font": {"size": Pt(12), "bold": True, "italic": True, "color": GRAY}},
                   {"paragraph_format": {"space_before": 0, "space_after": 0},
                    "font": {"size": Pt(11), "bold": True}}]
    },

    "project_summary": {
        "criterion": (cv_paragraphs_before_list, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                                         "space_before": Inches(.05)}}]
    },

    "lst_project_details": {
        "criterion": (cv_successive_list_elements, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                                         "space_before": Inches(.1), "space_after": 0},
                    "font": {"bold": True, "color": BLUE},
                    "is_list": True},
                   {"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                                         "space_before": 0, "space_after": 0},
                    "preserve_src_bold": True}]
    },

    "technical_environment": {
        "criterion": (cv_paragraphs_until_empty, {}),
        "styles": [{"paragraph_format": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                                         "left_indent": Inches(.5),
                                         "space_before": Inches(.1), "space_after": 0},
                    "font": {"bold": True, "color": ORANGE},
                    "empty_paragraphs_after": 1}]
    }
}


# MARK: -------- WRITE FUNCTIONS


def set_run_font(run: Run, font: dict) -> None:
    """
    Set `run`'s font attributes to the values in `font`.
    """
    for key, value in font.items():
        if key == "color":
            run.font.color.rgb = RGBColor.from_string(value)
        else:
            setattr(run.font, key, value)


def write_paragraph(src: Paragraph | _Cell, dst_p: Paragraph, style: dict) -> None:
    """
    Write `src`'s text into `dst_p` with the paragraph format and font specified in `style`.
    `src` can either be a paragraph or a cell. If it is a cell, the concatenation of the text of all
    its paragraphs is used.
    """
    print(f"\033[92m  Writing paragraph '{src.text}'\n  with style {style}\033[0m")

    # Paragraph format
    for key, value in style["paragraph_format"].items():
        setattr(dst_p.paragraph_format, key, value)

    # Font (this is also where the text is added to the paragraph)
    # Font name needs to be set to default value explicitly if not present
    style["font"].setdefault("name", DEFAULT_STYLE["font"]["name"])
    paragraphs = src.paragraphs if isinstance(src, _Cell) else [src]
    if style.get("preserve_src_bold", False):
        # The source paragraph must be processed run by run
        for src_p in (p for p in paragraphs if p.text):
            for src_run in (r for r in src_p.runs if r.text):
                dst_run = dst_p.add_run(src_run.text)
                set_run_font(dst_run, style["font"])
                dst_run.font.bold = src_run.font.bold
    else:
        src_text = "\n".join([p.text for p in paragraphs if p.text])
        dst_run = dst_p.add_run(src_text)
        set_run_font(dst_run, style["font"])


def write_table(src: Table | Paragraph, dst_table: Table, styles: list[dict]) -> None:
    """
    Write `src`'s text into `dst_table` with the paragraph format and font specified in `styles`.
    Set specific `dst_table` table properties based on `styles[0]`'s extra attributes as well.
    """
    if (indent_value := styles[0].get("table_left_indent")) is not None:
        # set value of table left indent
        tbl_ind = {"tblInd": {"w": str(indent_value.twips), "type": "dxa"}}
        rec_add_xml_children(dst_table._tblPr, tbl_ind)

    if (col_widths := styles[0].get("col_widths")) is not None:
        table_width = TEXT_WIDTH - (indent_value or 0)
        for i_col, dst_col in enumerate(dst_table.columns):
            col_width = col_widths.get(i_col) or round((table_width - sum(col_widths.values()))
                                                       / (len(src.columns) - len(col_widths)))
            dst_col.width = dst_table.cell(0, i_col).width = col_width

    if (margin_value := styles[0].get("side_margins")) is not None:
        # set value of table left and right margins
        cell_mar_attrs = {"w": str(margin_value), "type": "dxa"}
        tbl_cell_mar = {"tblCellMar": {"left": cell_mar_attrs, "right": cell_mar_attrs}}
        rec_add_xml_children(dst_table._tblPr, tbl_cell_mar)

    if styles[0].get("bottom_border"):
        border_attrs = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
        tbl_borders = {"tblBorders": {"bottom": border_attrs}}
        rec_add_xml_children(dst_table._tblPr, tbl_borders)

    for i_row, src_row in enumerate(src.rows):
        i_cell = 0
        while i_cell < len(src_row.cells):
            src_cell = src_row.cells[i_cell]
            dst_p = dst_table.cell(i_row, i_cell).paragraphs[0]
            # one style per column if there are enough columns, else one style per cell
            i_style = (i_cell if len(styles) <= len(src.columns)
                       else i_row * len(src.columns) + i_cell)
            write_paragraph(src_cell, dst_p, styles[min(i_style, len(styles) - 1)])

            # look for merged cells in the current row (cells over multiple rows are not supported)
            if (span := src_cell.grid_span) is not None and span > 1:
                dst_table.cell(i_row, i_cell).merge(dst_table.cell(i_row, i_cell + span - 1))
                i_cell = i_cell + span - 1

            i_cell += 1


def write_block(block: list[DocElm], dst_dc: Document, styles: list[dict]) -> None:
    """
    Write a block returned by one of the criterion validators into the destination DC with the
    appropriate style.
    """
    if len(block) == 0:
        return

    for style in styles:
        # Styles paragraph_format and font must be set if not specified in the block definition
        style.setdefault("paragraph_format", {})
        style.setdefault("font", {})

    main_style = styles[0]
    first_ilvl = get_ilvl(block[0])
    second_ilvl = get_ilvl(block[1]) if len(block) > 1 else 1

    for e in block:

        if isinstance(e, Table):
            dst_table = dst_dc.add_table(len(e.rows), len(e.columns))
            write_table(e, dst_table, styles)

        elif isinstance(e, Paragraph):
            dst_p = dst_dc.add_paragraph()
            if main_style.get("is_list"):
                ilvl = 0 if get_ilvl(e) == first_ilvl else get_ilvl(e) + 1 - second_ilvl
                numPr = dst_p._p.get_or_add_pPr().get_or_add_numPr()
                numPr.get_or_add_ilvl().set(qn("w:val"), str(ilvl))
                numPr.get_or_add_numId().set(qn("w:val"), "1")
                write_paragraph(e, dst_p, styles[min(ilvl, len(styles) - 1)])
            else:
                write_paragraph(e, dst_p, styles[0])

    for _ in range(main_style.get("empty_paragraphs_after", 0)):
        dst_dc.add_paragraph()


def write_unexpected_block(block: list[DocElm], dst_dc: Document) -> None:
    """
    Write a block that was identified as unexpected by one of the criterion validators, into the
    destination DC, with a yellow highlight.
    """
    for e in block:
        if isinstance(e, Table):
            dst_table = dst_dc.add_table(len(e.rows), len(e.columns))
            for i_row, src_row in enumerate(e.rows):
                for i_cell, src_cell in enumerate(src_row.cells):
                    dst_cell = dst_table.rows[i_row].cells[i_cell]
                    dst_cell.text = src_cell.text
                    rec_add_xml_children(dst_cell._tc.tcPr, {"shd": {"fill": "FFFF00"}})
        elif isinstance(e, Paragraph):
            dst_dc.add_paragraph(e.text).runs[0].font.highlight_color = WD_COLOR.YELLOW


# MARK: -------- MAIN


def parse_and_reformat_dc(src_dc: Document, dst_path: str) -> None:
    """
    Generate a reformated version of `src_dc` and save it to `dst_path`.
    The reformat relies on a static analysis of the document: a sequencial partition defines the
    semantic sections of a DC and the format associated to each of them, then the content of the
    input document is distributed into these sections using criterion validating functions.
    """
    dst_dc = docx.Document(TEMPLATE_DC_PATH)

    # Use a tee iterator to be able to peek it
    # Using tee()'s default argument value n=2, see https://github.com/python/cpython/issues/126701
    _, it_src_dc = tee(src_dc.iter_inner_content())
    _BLOCK_NAMES = tuple(BLOCK_DEFINITIONS)

    i = 0
    while True:
        print(f"\033[96mCurrent block: {_BLOCK_NAMES[i]}\033[0m")
        block_definition = BLOCK_DEFINITIONS[_BLOCK_NAMES[i]]
        criterion_validator, kwargs = block_definition["criterion"]
        styles = block_definition["styles"]
        try:
            print(f"\033[93m  {criterion_validator.__name__}:\033[0m")
            expected_block, unexpected_block = criterion_validator(it_src_dc, **kwargs)
            print(f"\033[93m  {len(expected_block)=}, {len(unexpected_block)=}\033[0m")
            write_block(expected_block, dst_dc, styles)
            write_unexpected_block(unexpected_block, dst_dc)
        except StopIteration:
            print("  Reached EOF!")
            break

        i += 1
        if i == len(_BLOCK_NAMES):
            # When reaching the end of block definitions, loop back from tbl_company_header
            i = _BLOCK_NAMES.index("tbl_company_header")

    dst_dc.save(dst_path)


def main(src_path: str = "test/default.docx", dst_path: str | None = None) -> None:
    src_dc = docx.Document(src_path)
    parse_and_reformat_dc(src_dc, dst_path or src_path.replace(".docx", "_reformated.docx"))


if __name__ == "__main__":
    parser = ArgumentParser(description="Write a reformated version of the input .docx file "
                                        "based on a static analysis and predefined styles.")

    parser.add_argument("input_file", metavar="input-file", nargs="?", default="test/default.docx",
                        help="Path to the input .docx file.")

    parser.add_argument("-o", "--output-path", metavar="PATH",
                        help="Specify the path to write the output file to.")

    parser_args = parser.parse_args()
    parser_kwargs = {"dst_path": parser_args.output_path}

    main(parser_args.input_file, **parser_kwargs)
