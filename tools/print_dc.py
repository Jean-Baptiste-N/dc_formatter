#!/usr/bin/env python3

from argparse import ArgumentParser
import json
import xml.etree.ElementTree as ET
from xml.dom import minidom
from pathlib import Path

import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .utils import get_ilvl


# Below are defined lists of attributes for some python-docx objects.
# An "_" is prepended to non-literal attributes. Those also have their own attribute list.
# An "_" is appended to iterable attributes.
DOCUMENT_ATTRS = ["DOCUMENT",
                  "_inline_shapes_", "_paragraphs_", "_sections_", "_styles_", "_tables_"]
SECTION_ATTRS = ["SECTION", "bottom_margin", "different_first_page_header_footer",
                 "even_page_footer", "even_page_header", "first_page_footer", "first_page_header",
                 "footer", "footer_distance", "gutter", "header", "header_distance", "left_margin",
                 "orientation", "page_height", "page_width", "right_margin", "start_type",
                 "top_margin"]
INLINE_SHAPE_ATTRS = ["INLINE_SHAPE", "height", "type", "width"]
TABLE_ATTRS = ["TABLE", "alignment", "autofit", "table_direction", "_style", "_columns_", "_rows_"]
COLUMN_ATTRS = ["COLUMN", "width", "cells_"]
ROW_ATTRS = ["ROW", "grid_cols_after", "grid_cols_before", "height", "height_rule", "_cells_"]
CELL_ATTRS = ["CELL", "grid_span", "text", "vertical_alignment", "width",
              "_paragraphs_", "_tables_"]
PARAGRAPH_ATTRS = ["PARAGRAPH", "alignment", "contains_page_break", "text",
                   "hyperlinks_", "rendered_page_breaks_",
                   "_paragraph_format", "_style", "_runs_"]
PARAGRAPH_FORMAT_ATTRS = ["PARAGRAPH_FORMAT", "alignment", "first_line_indent", "keep_together",
                          "keep_with_next", "left_indent", "line_spacing", "line_spacing_rule",
                          "page_break_before", "right_indent", "space_after", "space_before",
                          "widow_control", "_tab_stops_"]
BASE_STYLE_ATTRS = ["BASE_STYLE", "builtin", "hidden", "locked", "name", "priority", "quick_style",
                    "style_id", "type", "unhide_when_used"]
CHAR_STYLE_ATTRS = ["CHAR_STYLE", "builtin", "hidden", "locked", "name", "priority", "quick_style",
                    "style_id", "type", "unhide_when_used", "_base_style", "_font"]
PARAGRAPH_STYLE_ATTRS = ["PARAGRAPH_STYLE", "builtin", "hidden", "locked", "name",
                         "next_paragraph_style", "priority", "quick_style", "style_id", "type",
                         "unhide_when_used", "_base_style", "_font", "_paragraph_format"]
TABLE_STYLE_ATTRS = ["TABLE_STYLE", "builtin", "hidden", "locked", "name", "next_paragraph_style",
                     "priority", "quick_style", "style_id", "type", "unhide_when_used",
                     "_base_style", "_font", "_paragraph_format"]
RUN_ATTRS = ["RUN", "bold", "contains_page_break", "italic", "text", "underline", "_font", "_style"]
FONT_ATTRS = ["FONT", "all_caps", "bold", "complex_script", "cs_bold", "cs_italic", "double_strike",
              "emboss", "hidden", "highlight_color", "imprint", "italic", "math", "name",
              "no_proof", "outline", "rtl", "shadow", "size", "small_caps", "snap_to_grid",
              "spec_vanish", "strike", "subscript", "superscript", "underline", "web_hidden",
              "_color"]
TAB_STOP_ATTRS = ["TABSTOP", "alignment", "leader", "position"]
COLOR_ATTRS = ["COLOR", "rgb", "theme_color", "type"]

# This dict maps some non-literal attributes from the lists above, to their own attribute list.
# It is necessary to disambiguate some homonymous attributes from different objects (e.g. "style").
# A key is the aatribute name prefixed by the name of the list it is from ("ATTRS" suffix cut off).
ATTR_LIST_MAP = {"DOCUMENT_styles_": CHAR_STYLE_ATTRS,
                 "TABLE_style": TABLE_STYLE_ATTRS,
                 "PARAGRAPH_style": PARAGRAPH_STYLE_ATTRS,
                 "RUN_style": CHAR_STYLE_ATTRS}


def c(code: int | str) -> str:
    """
    Macro returning the ANSi Select Graphic Rendition containing `code`.
    """
    return f"\033[{code}m"


RESET, WHITE, L_GREY, D_GREY, BLACK, PINK, GREEN = c(0), c(97), c(37), c(90), c(30), c(91), c(92)
YELLOW, ORANGE, BLUE, MAGENTA, PURPLE, TURQUOISE = c(93), c(33), c(94), c(95), c(35), c(36)
BG_GREEN, BG_YELLOW, BG_ORANGE = c("48;2;0;47;0"), c("48;2;31;31;0"), c("48;2;31;15;0")


def _get_child_attr_list(parent_attr_list: list[str], attr_name: str) -> list[str]:
    """
    Considering `attr_name` is from `parent_attr_list`, return the list of `attr_name`'s child
    attributes based on `attr_name` and `attr_list` if necessary.
    """
    try:
        return globals()[attr_name[1:].rstrip("s_").upper() + "_ATTRS"]
    except KeyError:
        return ATTR_LIST_MAP[f"{parent_attr_list[0]}{attr_name}"]


def triage_attrs(o, attr_names: list[str]) -> tuple[list[str], list[str], list[str], list[str]]:
    """
    Triage a list of attribute names based on the attribute values. Return 4 lists in this order:
    1. list of attributes whose value is None;
    2. list of attributes whose value is True;
    3. list of attributes whose value is False;
    4. list of remaining attributes.
    """
    nones, trues, falses, rest = [], [], [], []
    for attr_name in attr_names:
        if (attr_value := getattr(o, attr_name.strip("_"))) is None:
            nones.append(attr_name)
        elif attr_value is True:
            trues.append(attr_name)
        elif attr_value is False:
            falses.append(attr_name)
        else:
            rest.append(attr_name)
    return nones, trues, falses, rest


def rec_is_empty(o, attr_list: list[str]) -> bool:
    """
    Return True if the object `o` only contains attributes that are:
    - literals equal to `None`, or
    - iterables of length 0, or
    - non-literal attributes with only empty sub-attributes, or
    - iterables of non-literals with only empty sub-attributes.
    """
    _, trues, falses, rest = triage_attrs(o, attr_list[1:])
    if trues or falses:
        return False

    for attr_name in rest:
        attr_value = getattr(o, attr_name.strip("_"))
        if attr_name[0] == attr_name[-1] == "_":
            # attribute is an iterable of non-literals: any element not empty => parent not empty
            for e in attr_value:
                child_attr_list = _get_child_attr_list(attr_list, attr_name)
                if not rec_is_empty(e, child_attr_list):
                    return False
        elif attr_name.startswith("_"):
            # attribute is a non-literal: not empty => parent not empty
            child_attr_list = _get_child_attr_list(attr_list, attr_name)
            if not rec_is_empty(attr_value, child_attr_list):
                return False
        elif attr_name.endswith("_"):
            # attribute is an iterable of literals: not of length 0 => parent not empty
            if len(attr_value) != 0:
                return False
        else:
            # attribute is a literal: parent not empty
            return False
    return True


def print_attrs(o) -> None:
    """
    Traverse the attributes of the object `o` and print them as a tree hierarchy.
    """
    INDENT = "    "

    def rec_print_attrs(o, attr_list: list[str], prefix: str) -> None:
        if rec_is_empty(o, attr_list):
            print(f"{prefix}{PURPLE}(empty){RESET}")
        else:
            print(f"{prefix}{TURQUOISE}non empty{RESET}")
        nones, trues, falses, rest = triage_attrs(o, attr_list[1:])
        print(f"{prefix}{D_GREY}Nones = {nones}{RESET}\n" if nones else "", end="")
        print(f"{prefix}{PINK}Trues = {trues}{RESET}\n" if trues else "", end="")
        print(f"{prefix}{BLUE}Falses = {falses}{RESET}\n" if falses else "", end="")

        for attr_name in rest:
            attr_value = getattr(o, attr_name.strip("_"))
            if attr_name[0] == attr_name[-1] == "_":
                # attribute is an iterable of non-literals
                print(f"{prefix}{attr_name[1:-1]} ({len(attr_value)} elms)")
                new_prefix = prefix + INDENT
                for i, e in enumerate(attr_value):
                    print(f"{new_prefix}{attr_name[1:-2]}[{i}]:")
                    child_attr_list = _get_child_attr_list(attr_list, attr_name)
                    rec_print_attrs(e, child_attr_list, new_prefix + INDENT)
            elif attr_name.startswith("_"):
                # attribute is a non-literal
                print(f"{prefix}{attr_name[1:]}:")
                child_attr_list = _get_child_attr_list(attr_list, attr_name)
                rec_print_attrs(attr_value, child_attr_list, prefix + INDENT)
            elif attr_name.endswith("_"):
                # attribute is an iterable of literals (this case should not occur)
                print(f"{prefix}{attr_name[:-1]} ({len(attr_value)} elms) = '{attr_value}'")
            else:
                # attribute is a literal
                print(f"{prefix}{attr_name} = '{attr_value}'")

    print(f"{type(o)=}")
    if isinstance(o, Table):
        rec_print_attrs(o, TABLE_ATTRS, "")
    elif isinstance(o, Paragraph):
        rec_print_attrs(o, PARAGRAPH_ATTRS, "")
    elif isinstance(o, Run):
        rec_print_attrs(o, RUN_ATTRS, "")


def get_font_children(o, prefix: str | None = None) -> dict[str, dict]:
    """
    Return a mapping of every non-empty `font` child attribute of `o` with its location in `o`.

    N.B. non-empty child does not mean that ALL attributes are non-empty.
    """
    non_empty_font_children = {}

    def rec_get_font_children(o, attr_list: list[str], prefix: str) -> None:
        _, _, _, rest = triage_attrs(o, attr_list[1:])

        for attr_name in rest:
            attr_value = getattr(o, attr_name.strip("_"))
            new_prefix = f"{prefix}.{attr_name.strip('_')}"
            if attr_name[0] == attr_name[-1] == "_":
                # attribute is an iterable of non-literals
                for i, e in enumerate(attr_value):
                    child_attr_list = _get_child_attr_list(attr_list, attr_name)
                    rec_get_font_children(e, child_attr_list, f"{new_prefix}[{i}]")
            elif attr_name.startswith("_"):
                # attribute is a non-literal
                child_attr_list = _get_child_attr_list(attr_list, attr_name)
                if attr_name == "_font" and not rec_is_empty(attr_value, child_attr_list):
                    color = None if attr_value.color.rgb is None else str(attr_value.color.rgb)
                    font_props = {"name": attr_value.name, "size": attr_value.size, "color": color,
                                  "bold": attr_value.bold, "italic": attr_value.italic,
                                  "underline": attr_value.underline}
                    non_empty_font_children[new_prefix] = font_props
                rec_get_font_children(attr_value, child_attr_list, new_prefix)

    if isinstance(o, Table):
        rec_get_font_children(o, TABLE_ATTRS, prefix or "t")
    elif isinstance(o, Paragraph):
        rec_get_font_children(o, PARAGRAPH_ATTRS, prefix or "p")
    elif isinstance(o, Run):
        rec_get_font_children(o, RUN_ATTRS, prefix or "r")
    return non_empty_font_children


def get_font_stats(o, prefix: str = "p") -> str:
    """
    Return a string giving statistics about `font` attributes contained in `o`:
    - total number of `font` attributes,
    - number of `font` attributes located in `run[x]` attributes,
    - number of `font` attributes located in `run[x].style` attributes,
    - all fonts defined as unique based on their font name, size and color.
    """
    fonts = get_font_children(o, prefix)
    unique_fonts = {}
    npsf = nrf = nrsf = 0
    for k, v in fonts.items():
        if v not in unique_fonts.values():
            unique_fonts[k] = v
        if "runs[" in k and k.endswith("].font"):
            nrf += 1
        elif "runs[" in k and k.endswith("].style.font"):
            nrsf += 1
        else:
            npsf += 1
    nf, nuf = len(fonts), len(unique_fonts)

    # colorized font stat strings
    str_nf = f"{WHITE if nf > 0 else D_GREY}{nf:3}"
    str_nuf = f"{MAGENTA if nuf > 1 else D_GREY}{nuf:3}"
    str_npsf = f"{L_GREY if npsf > 0 else D_GREY}{npsf:3}"
    str_nrf = f"{L_GREY if nrf > 0 else D_GREY}{nrf:3}"
    str_nrsf = f"{L_GREY if nrsf > 0 else D_GREY}{nrsf:3}"

    c_uf = GREEN if prefix == "t" else YELLOW if prefix.startswith(("p", "rw")) else ORANGE

    def str_uf(uf: dict) -> str:
        s = ""
        for prop in ("size", "color"):
            s += f"{uf[prop]} " if uf[prop] else f"{D_GREY}-{c_uf} "
        for prop in ("bold", "italic", "underline"):
            s += f"{WHITE if uf[prop] else D_GREY if uf[prop] is False else BLACK}{prop[0].upper()}"
        return f"{c_uf}{s}{c_uf}"

    str_unique_fonts = " - ".join([f"{k}: [{str_uf(v)}]" for k, v in unique_fonts.items()])
    return f"{str_nf}{str_nuf}{str_npsf}{str_nrf}{str_nrsf} {c_uf}{str_unique_fonts}"


def print_content_with_font(dc: Document) -> None:
    """
    Print all paragraphs of the document located at `path`, with font statistics prepended to each
    paragraph.
    """
    sl = [f"S[{i}]: {len(list(s.iter_inner_content()))} elms" for i, s in enumerate(dc.sections)]
    print(f"{MAGENTA}{' - '.join(sl)}{RESET}")

    def _print_paragraph_content_with_font(p: Paragraph, line_pfx: str, font_pfx: str) -> None:
        in_table = font_pfx.startswith(("t", "rw"))
        str_ilvl = f"{ilvl=} | " if (ilvl := get_ilvl(p)) is not None else ""
        # print(f"{line_pfx}{get_font_stats(p, font_pfx)} | {p.style.style_id} |{RESET} '{p.text}'")
        print(f"{line_pfx} | {p.style.style_id:30} | {str_ilvl}{WHITE}'{p.text}'{RESET}")
        if len(p.runs) > 1:
            for j, r in enumerate(p.runs):
                line_pfx = f"{BG_ORANGE}{ORANGE}" + (f"  r{j:<7}" if in_table else f" r{j:<7}")
                font_pfx = f"runs[{j}]"
                # print(f"{line_pfx}{get_font_stats(r, font_pfx)} |{WHITE} '{r.text}'{RESET}")
                print(f"{line_pfx} | {r.style.style_id:30}{WHITE}'{r.text}'{RESET}")

    for i, e in enumerate(dc.iter_inner_content()):
        if isinstance(e, Table):
            print(f"{BG_GREEN}{GREEN}{i:<8}{get_font_stats(e, 't')}{RESET}")
            for ir, r in enumerate(e.rows):
                for ic, cell in enumerate(r.cells):
                    for ip, p in enumerate(cell.paragraphs):
                        line_pfx = f"{BG_YELLOW}{YELLOW} rw{ir}cl{ic}p{ip}"
                        font_pfx = f"rw[{ir}].cl[{ic}].p[{ip}]"
                        _print_paragraph_content_with_font(p, line_pfx, font_pfx)

        else:  # Paragraph
            _print_paragraph_content_with_font(e, f"{BG_YELLOW}{YELLOW}{i:<8}", "p")



def build_document_structure(dc: Document) -> dict:
    """
    Build a hierarchical JSON representation of the document structure.
    Includes section information and page/section breaks.
    """
    structure = {
        "document": {
            "type": "Document",
            "sections_count": len(dc.sections),
            "sections": []
        }
    }
    
    # Ajouter les informations des sections si multiples
    if len(dc.sections) > 1:
        for i, section in enumerate(dc.sections):
            section_info = {
                "id": i,
                "start_type": str(section.start_type)
            }
            structure["document"]["sections"].append(section_info)
    
    structure["document"]["content"] = []
    
    for i, element in enumerate(dc.iter_inner_content()):
        if isinstance(element, Table):
            table_struct = {
                "id": i,
                "type": "Table",
                "rows": len(element.rows),
                "cols": len(element.columns),
                "rows_data": []
            }
            for row_idx, row in enumerate(element.rows):
                row_data = {
                    "row_id": row_idx,
                    "cells": []
                }
                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "cell_id": cell_idx,
                        "grid_span": cell.grid_span,
                        "text": cell.text,
                        "paragraphs": []
                    }
                    for para in cell.paragraphs:
                        para_struct = _build_paragraph_structure(para)
                        cell_data["paragraphs"].append(para_struct)
                    row_data["cells"].append(cell_data)
                table_struct["rows_data"].append(row_data)
            structure["document"]["content"].append(table_struct)
        
        elif isinstance(element, Paragraph):
            para_struct = _build_paragraph_structure(element, include_runs=True)
            structure["document"]["content"].append(para_struct)
    
    return structure


def _build_paragraph_structure(p: Paragraph, include_runs: bool = True) -> dict:
    """
    Build a hierarchical structure for a single paragraph.
    Includes page breaks, section breaks, and section breaks.
    """
    ilvl = get_ilvl(p)
    para_struct = {
        "type": "Paragraph",
        "text": p.text,
        "style_id": p.style.style_id if p.style else None,
        "alignment": str(p.alignment) if p.alignment is not None else None,
        "ilvl": ilvl,
        "page_break_before": p.paragraph_format.page_break_before,
        "keep_with_next": p.paragraph_format.keep_with_next,
        "keep_together": p.paragraph_format.keep_together,
        "section_break": False  # Default
    }
    
    # Détection de saut de section
    try:
        if p._element.pPr is not None and p._element.pPr.sectPr is not None:
            para_struct["section_break"] = True
    except (AttributeError, TypeError):
        pass
    
    if include_runs and p.runs:
        para_struct["runs"] = []
        for run_idx, run in enumerate(p.runs):
            run_struct = {
                "run_id": run_idx,
                "text": run.text,
                "font": {
                    "name": run.font.name,
                    "size": str(run.font.size) if run.font.size else None,
                    "bold": run.font.bold,
                    "italic": run.font.italic,
                    "underline": run.font.underline,
                    "color": str(run.font.color.rgb) if run.font.color.rgb else None
                }
            }
            para_struct["runs"].append(run_struct)
    
    return para_struct


def build_document_xml(dc: Document) -> ET.Element:
    """
    Build an XML tree representation of the document structure with line numbers.
    Includes page breaks, section breaks, and section information.
    """
    root = ET.Element("document")
    root.set("sections", str(len(dc.sections)))
    
    # Ajouter les sections du document
    if len(dc.sections) > 1:
        sections_elem = ET.SubElement(root, "sections_info")
        for i, section in enumerate(dc.sections):
            section_elem = ET.SubElement(sections_elem, "section")
            section_elem.set("id", str(i))
            section_elem.set("start_type", str(section.start_type))
    
    line_num = 0
    for element in dc.iter_inner_content():
        if isinstance(element, Table):
            table_elem = ET.SubElement(root, "table")
            table_elem.set("line", str(line_num))
            table_elem.set("rows", str(len(element.rows)))
            table_elem.set("cols", str(len(element.columns)))
            
            for row_idx, row in enumerate(element.rows):
                row_elem = ET.SubElement(table_elem, "row")
                row_elem.set("id", str(row_idx))
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_elem = ET.SubElement(row_elem, "cell")
                    cell_elem.set("id", str(cell_idx))
                    if cell.grid_span and cell.grid_span > 1:
                        cell_elem.set("grid_span", str(cell.grid_span))
                    
                    for para in cell.paragraphs:
                        para_elem = _build_paragraph_xml(para, line_num)
                        cell_elem.append(para_elem)
                        line_num += 1
            
            line_num += 1
        
        elif isinstance(element, Paragraph):
            para_elem = _build_paragraph_xml(element, line_num)
            root.append(para_elem)
            line_num += 1
    
    return root


def _build_paragraph_xml(p: Paragraph, line_num: int) -> ET.Element:
    """
    Build an XML element for a single paragraph with runs and font details.
    Includes page breaks, section breaks, and section breaks.
    """
    para_elem = ET.Element("paragraph")
    para_elem.set("line", str(line_num))
    para_elem.set("style", p.style.style_id if p.style else "None")
    
    if p.alignment is not None:
        para_elem.set("alignment", str(p.alignment))
    
    ilvl = get_ilvl(p)
    if ilvl is not None:
        para_elem.set("ilvl", str(ilvl))
    
    # Sauts de page et de section
    if p.paragraph_format.page_break_before:
        para_elem.set("page_break_before", "True")
    if p.paragraph_format.keep_with_next:
        para_elem.set("keep_with_next", "True")
    if p.paragraph_format.keep_together:
        para_elem.set("keep_together", "True")
    
    # Détection de saut de section (section properties dans le XML interne)
    try:
        if p._element.pPr is not None and p._element.pPr.sectPr is not None:
            para_elem.set("section_break", "True")
    except (AttributeError, TypeError):
        pass
    
    # Contenu texte
    text_elem = ET.SubElement(para_elem, "text")
    text_elem.text = p.text if p.text else ""
    
    # Détail des runs
    if p.runs:
        runs_elem = ET.SubElement(para_elem, "runs")
        for run_idx, run in enumerate(p.runs):
            run_elem = ET.SubElement(runs_elem, "run")
            run_elem.set("id", str(run_idx))
            run_elem.text = run.text if run.text else ""
            
            # Propriétés de police
            font_elem = ET.SubElement(run_elem, "font")
            if run.font.name:
                font_elem.set("name", run.font.name)
            if run.font.size:
                font_elem.set("size", str(run.font.size))
            if run.font.bold:
                font_elem.set("bold", str(run.font.bold))
            if run.font.italic:
                font_elem.set("italic", str(run.font.italic))
            if run.font.underline:
                font_elem.set("underline", str(run.font.underline))
            if run.font.color.rgb:
                font_elem.set("color", str(run.font.color.rgb))
    
    return para_elem


def _prettify_and_save_xml(elem: ET.Element, output_path: str) -> None:
    """
    Pretty-print XML tree and save it to file.
    """
    rough_string = ET.tostring(elem, encoding='utf-8')
    reparsed = minidom.parseString(rough_string)
    pretty_xml = reparsed.toprettyxml(indent="  ")
    
    # Remove XML declaration line if present
    lines = pretty_xml.split('\n')
    if lines[0].startswith('<?xml'):
        lines = lines[1:]
    pretty_xml = '\n'.join(lines)
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)


def main(path: str = "test/default.docx", print_object: str = None, export_json: bool = False, export_xml: bool = False) -> None:
    dc = docx.Document(path)
    
    # Only print to terminal if no export is specified
    if not export_json and not export_xml:
        if print_object is not None:
            # print_attrs([e for e in dc.iter_inner_content()][70])
            print_attrs(eval(print_object))
        else:
            print_content_with_font(dc)
    
    # Créer le dossier structures s'il n'existe pas
    structures_dir = Path("structures")
    if export_json or export_xml:
        structures_dir.mkdir(exist_ok=True)
    
    if export_json:
        structure = build_document_structure(dc)
        output_path = structures_dir / (Path(path).stem + "_structure.json")
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(structure, f, indent=2, ensure_ascii=False)
        print(f"✓ Structure exported to {output_path}")
    
    if export_xml:
        root = build_document_xml(dc)
        output_path = structures_dir / (Path(path).stem + "_structure.xml")
        _prettify_and_save_xml(root, output_path)
        print(f"✓ Structure exported to {output_path}")

if __name__ == "__main__":
    parser = ArgumentParser(description="Print text extracted from input .docx file.")

    parser.add_argument("path", nargs="?", default="test/default.docx", help="Path to .docx file.")

    parser.add_argument("--print", dest="print_object", metavar="OBJECT",
                        help="Print an attribute hierarchy tree of %(metavar)s.")
    
    parser.add_argument("--export_json", action="store_true",
                        help="Export document structure as {name}_structure.json")
    
    parser.add_argument("--export_xml", action="store_true",
                        help="Export document structure as {name}_structure.xml")

    parser_args = parser.parse_args()
    parser_kwargs = {
        "print_object": parser_args.print_object,
        "export_json": parser_args.export_json,
        "export_xml": parser_args.export_xml
    }

    main(parser_args.path, **parser_kwargs)
