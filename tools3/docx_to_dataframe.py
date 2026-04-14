"""
Script pour extraire le texte d'un fichier DOCX et le reporter dans un dataframe pandas.
Support de plusieurs formats de docx et extraction structurée du contenu.
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
import zipfile
from typing import List, Dict, Tuple
import xml.etree.ElementTree as ET


def extract_text_from_docx(docx_file: str) -> Tuple[List[Dict], str]:
    """
    Extrait le texte d'un fichier DOCX avec structure et métadonnées.
    
    Args:
        docx_file (str): Chemin vers le fichier .docx
        
    Returns:
        Tuple[List[Dict], str]: (liste de dictionnaires avec le texte, nom du fichier)
    """
    data = []
    
    try:
        # Charger le document
        doc = Document(docx_file)
        
        # Extraire le texte par paragraphe
        for para_idx, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            
            if text:  # Ne garder que les paragraphes non-vides
                try:
                    level = paragraph.paragraph_format.outline_level
                    if level is None:
                        level = 0
                except:
                    level = 0
                    
                data.append({
                    'paragraph_id': para_idx,
                    'text': text,
                    'style': paragraph.style.name,
                    'level': level,
                    'type': 'paragraph'
                })
        
        # Extraire le texte des tableaux
        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows, 1):
                for cell_idx, cell in enumerate(row.cells, 1):
                    cell_text = cell.text.strip()
                    if cell_text:
                        data.append({
                            'paragraph_id': f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                            'text': cell_text,
                            'style': 'table_cell',
                            'level': None,
                            'type': 'table'
                        })
        
        print(f"✓ {len(data)} éléments texte extraits de {Path(docx_file).name}")
        
    except Exception as e:
        print(f"✗ Erreur lors de la lecture de {docx_file}: {e}")
        return [], Path(docx_file).name
    
    return data, Path(docx_file).name


def extract_text_raw_xml(docx_file: str) -> Tuple[List[Dict], str]:
    """
    Alternative: extrait le texte en parsant le XML brut du document.xml.
    Utile pour les fichiers corrompus ou mal formatés.
    
    Args:
        docx_file (str): Chemin vers le fichier .docx
        
    Returns:
        Tuple[List[Dict], str]: (liste de dictionnaires avec le texte, nom du fichier)
    """
    data = []
    
    try:
        with zipfile.ZipFile(docx_file, 'r') as zip_ref:
            # Lire le fichier document.xml
            xml_content = zip_ref.read('word/document.xml').decode('utf-8')
            
            # Parser le XML
            root = ET.fromstring(xml_content)
            
            # Namespace pour Word
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Extraire tous les paragraphes
            paragraphs = root.findall('.//w:p', ns)
            
            for para_idx, paragraph in enumerate(paragraphs, 1):
                # Extraire tous les textes du paragraphe
                texts = []
                for text_elem in paragraph.findall('.//w:t', ns):
                    if text_elem.text:
                        texts.append(text_elem.text)
                
                full_text = ''.join(texts).strip()
                
                if full_text:  # Ne garder que les paragraphes non-vides
                    data.append({
                        'paragraph_id': para_idx,
                        'text': full_text,
                        'style': None,
                        'level': None,
                        'type': 'paragraph'
                    })
        
        print(f"✓ {len(data)} paragraphes extraits de {Path(docx_file).name} (XML raw)")
        
    except Exception as e:
        print(f"✗ Erreur lors de la lecture XML de {docx_file}: {e}")
        return [], Path(docx_file).name
    
    return data, Path(docx_file).name


def docx_to_dataframe(docx_file: str, use_raw_xml: bool = False) -> pd.DataFrame:
    """
    Convertit un fichier DOCX en dataframe pandas.
    
    Args:
        docx_file (str): Chemin vers le fichier .docx
        use_raw_xml (bool): Si True, utilise le parsing XML brut (plus robuste)
                           Si False, utilise python-docx (mieux structuré)
        
    Returns:
        pd.DataFrame: Dataframe avec le contenu du DOCX
    """
    if use_raw_xml:
        data, filename = extract_text_raw_xml(docx_file)
    else:
        data, filename = extract_text_from_docx(docx_file)
    
    if not data:
        print(f"Avertissement: Aucun texte trouvé dans {filename}")
        return pd.DataFrame()
    
    df = pd.DataFrame(data)
    df['source_file'] = filename
    
    return df


def process_multiple_docx(docx_files: List[str], use_raw_xml: bool = False) -> pd.DataFrame:
    """
    Traite plusieurs fichiers DOCX et combine tous les textes dans un dataframe.
    
    Args:
        docx_files (List[str]): Liste des chemins vers les fichiers .docx
        use_raw_xml (bool): Si True, utilise le parsing XML brut
        
    Returns:
        pd.DataFrame: Dataframe combiné avec tout le contenu
    """
    all_data = []
    
    for docx_file in docx_files:
        df = docx_to_dataframe(docx_file, use_raw_xml=use_raw_xml)
        if not df.empty:
            all_data.append(df)
    
    if not all_data:
        print("Avertissement: Aucun texte trouvé dans aucun fichier")
        return pd.DataFrame()
    
    combined_df = pd.concat(all_data, ignore_index=True)
    return combined_df


if __name__ == "__main__":
    import sys
    
    # Exemple d'utilisation
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        use_xml = len(sys.argv) > 2 and sys.argv[2] == '--raw-xml'
        
        # Vérifier si c'est un fichier ou un répertoire
        path = Path(docx_path)
        
        if path.is_file() and path.suffix.lower() == '.docx':
            # Fichier unique
            df = docx_to_dataframe(str(path), use_raw_xml=use_xml)
            
            print(f"\n📊 Dataframe créé avec {len(df)} lignes et {len(df.columns)} colonnes")
            print("\nColonnes:", df.columns.tolist())
            print("\nPremiers résultats:")
            print(df.head(10))
            
            # Option: Sauvegarder en CSV
            csv_path = path.with_suffix('.csv')
            df.to_csv(csv_path, index=False, encoding='utf-8')
            print(f"\n✓ CSV sauvegardé: {csv_path}")
            
        elif path.is_dir():
            # Répertoire de fichiers DOCX
            docx_files = list(path.glob('*.docx'))
            if docx_files:
                print(f"Traitement de {len(docx_files)} fichier(s) DOCX...")
                df = process_multiple_docx([str(f) for f in docx_files], use_raw_xml=use_xml)
                
                print(f"\n📊 Dataframe combiné créé avec {len(df)} lignes et {len(df.columns)} colonnes")
                print("\nRésumé par fichier source:")
                print(df.groupby('source_file').size())
                
                # Sauvegarder en CSV
                csv_path = path / 'combined_output.csv'
                df.to_csv(csv_path, index=False, encoding='utf-8')
                print(f"\n✓ CSV combiné sauvegardé: {csv_path}")
            else:
                print(f"Aucun fichier .docx trouvé dans {path}")
        else:
            print(f"Erreur: {docx_path} n'est pas un fichier .docx valide")
    else:
        print("Usage:")
        print("  python docx_to_dataframe.py <chemin_vers_fichier.docx> [--raw-xml]")
        print("  python docx_to_dataframe.py <dossier_contenant_docx> [--raw-xml]")
        print("\nOptions:")
        print("  --raw-xml: Utilise le parsing XML brut (plus robuste pour les fichiers corrompus)")
