"""
Script pour injecter le contenu d'un JSON (balises) dans un template DOCX.
Entrée: fichier JSON avec paragraphes et tableaux
Sortie: DOCX dans le dossier renders/
"""

import json
from pathlib import Path
from argparse import ArgumentParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# # Approche template (2 lignes)
# template_doc = Document('assets/TEMPLATE.docx')

# KEYWORDS_HEADER_DOCUMENT = ["dossier de compétences", "dossier de competence", "dossier de competences", "dossier de competences"]
# KEYWORDS_MAIN_SKILLS = ["domaine de compétence", "domaine de competence", "domaines de compétence", "domaines de competence", "compétences principales", "competences principales", "compétence", "competence"]
# KEYWORDS_EDUCATION = ["formation", "diplôme", "diplome", "certification", "langue", "langues", "certifications", "diplômes", "diplomes"]
# KEYWORDS_PROFESSIONAL_EXPERIENCE = ["expérience professionnelle", "experience professionnelle", "expériences professionnelles", "experience professionnelles"]


def parse_alignment(align_str: str):
    """Convertit string d'alignment en WD_ALIGN_PARAGRAPH"""
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return align_map.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)


def add_paragraph_from_json(doc: Document, para_data: dict):
    """Ajoute un paragraphe au document à partir des données JSON"""
    para = doc.add_paragraph()

    # Appliquer les propriétés
    props = para_data.get('properties', {})

    # Alignment
    if 'alignment' in props:
        para.alignment = parse_alignment(props['alignment'])

    # Style (accéder à l'objet Style directement pour éviter le deprecated warning)
    if 'style' in props:
        try:
            style_name = props['style']
            # Obtenir l'objet Style du document
            if style_name in doc.styles:
                para.style = doc.styles[style_name]  # Utiliser l'objet Style, pas le nom
            else:
                # Essayer d'assigner par nom si l'objet n'existe pas
                para.style = style_name
        except (KeyError, KeyError):
            pass  # Ignorer si le style n'existe pas

    # Ajouter les runs
    if 'runs' in para_data:
        for run_data in para_data['runs']:
            # Vérifier si c'est un saut de page
            if run_data.get('page_break'):
                run = para.add_run()
                run.add_break(WD_BREAK_TYPE.PAGE)  # Saut de page avec type
                continue

            run_text = run_data.get('text', '')
            run = para.add_run(run_text)

            # Appliquer les propriétés du run
            run_props = run_data.get('properties', {})

            if run_props.get('bold'):
                run.bold = True
            if run_props.get('italic'):
                run.italic = True

            # Size (en points)
            if 'size' in run_props:
                try:
                    size_half_pt = int(run_props['size'])
                    run.font.size = Pt(size_half_pt / 2)  # XML Word utilise demi-points
                except:
                    pass

            # Color
            if 'color' in run_props:
                try:
                    color_hex = run_props['color']
                    run.font.color.rgb = RGBColor(
                        int(color_hex[0:2], 16),
                        int(color_hex[2:4], 16),
                        int(color_hex[4:6], 16)
                    )
                except:
                    pass

            # Font
            if 'font' in run_props:
                run.font.name = run_props['font']

    # Section break (saut de section) - ajouter via XML
    if 'section_break' in props:
        section_type = props['section_break']
        pPr = para._element.get_or_add_pPr()

        # Supprimer l'ancien sectPr s'il existe
        old_sectPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
        if old_sectPr is not None:
            pPr.remove(old_sectPr)

        # Créer un nouveau sectPr
        sectPr_xml = f'''
            <w:sectPr {nsdecls('w')}>
                <w:type w:val="{section_type}"/>
                <w:pgSz w:w="11920" w:h="16840" w:orient="portrait"/>
                <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720"/>
            </w:sectPr>
        '''.strip()
        sectPr_elem = parse_xml(sectPr_xml)
        pPr.append(sectPr_elem)

    # Texte simple si pas de runs
    elif 'text' in para_data:
        para.add_run(para_data['text'])


def set_table_borders(table, borders_data):
    """
    Applique les bordures à une table Word.

    borders_data: {
        'top': {'size': '12', 'color': '000000'} or None,
        'bottom': {'size': '12', 'color': '000000'} or None,
        'left': ..., 'right': ..., 'insideH': ..., 'insideV': ...
    }
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    # Accéder à la propriété tbl de la table
    tbl = table._element
    tblPr = tbl.tblPr

    # Créer/récupérer les propriétés de bordure (tblBorders)
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Construire le XML des bordures
    borders_xml = f'<w:tblBorders {nsdecls("w")}>'

    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_info = borders_data.get(side)
        if border_info is None:
            # Pas de bordure
            borders_xml += f'<w:{side} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        else:
            size = border_info.get('size', '12')
            color = border_info.get('color', '000000')
            borders_xml += f'<w:{side} w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'

    borders_xml += '</w:tblBorders>'

    tblBorders_elem = parse_xml(borders_xml)
    tblPr.append(tblBorders_elem)


def add_table_from_json(doc: Document, table_data: dict):
    """Ajoute un tableau au document à partir des données JSON"""
    rows = table_data.get('rows', [])
    if not rows:
        return

    # Nombre de colonnes
    col_count = table_data.get('col_count', len(rows[0]['cells']) if rows else 0)

    # Créer le tableau
    table = doc.add_table(rows=len(rows), cols=col_count)

    # Note: Le style de table n'est pas assigné ici car les styles personnalisés
    # sont gérés par le template DOCX. Seules les propriétés structurelles (largeur,
    # hauteur, bordures) sont appliquées à partir du JSON.

    # Appliquer les propriétés du tableau
    props = table_data.get('properties', {})

    # Appliquer les bordures si définies dans les propriétés
    borders_data = props.get('borders')
    if borders_data:
        set_table_borders(table, borders_data)

    if 'table_width' in props:
        # Appliquer la largeur du tableau (en twips)
        # Conversion: 1 twip = 2.54 cm / 1440 = 0.00176389 cm
        try:
            width_twips = int(props['table_width'])
            width_cm = (width_twips * 2.54) / 1440
            table.width = Cm(width_cm)
        except (ValueError, TypeError):
            pass

    # Remplir les cellules
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]

        # Appliquer la hauteur de la ligne si elle existe
        if 'height' in row_data:
            try:
                height_twips = int(row_data['height'])
                height_cm = (height_twips * 2.54) / 1440
                row.height = Cm(height_cm)
            except (ValueError, TypeError):
                pass

        for col_idx, cell_data in enumerate(row_data.get('cells', [])):
            cell = row.cells[col_idx]

            # Appliquer l'alignement de la cellule (horizontal et vertical)
            cell_props = cell_data.get('properties', {})
            if 'vAlign' in cell_props:
                # Alignement vertical : 'top', 'center', 'bottom'
                v_align = cell_props['vAlign'].lower()
                if v_align == 'center':
                    cell.vertical_alignment = 1  # WD_ALIGN_VERTICAL.CENTER
                elif v_align == 'bottom':
                    cell.vertical_alignment = 2  # WD_ALIGN_VERTICAL.BOTTOM
                else:  # top
                    cell.vertical_alignment = 0  # WD_ALIGN_VERTICAL.TOP

            # Appliquer la largeur de la cellule si elle existe
            if 'width' in cell_data:
                try:
                    width_twips = int(cell_data['width'])
                    width_cm = (width_twips * 2.54) / 1440
                    cell.width = Cm(width_cm)
                except (ValueError, TypeError):
                    pass

            # Ajouter les paragraphes de la cellule
            paragraphs_list = cell_data.get('paragraphs', [])

            if paragraphs_list:
                # Réutiliser le premier paragraphe existant pour le premier contenu
                para = cell.paragraphs[0]
                para_data = paragraphs_list[0]

                # Appliquer les propriétés du paragraphe
                para_props = para_data.get('properties', {})

                # Style
                if 'style' in para_props:
                    try:
                        if para_props['style'] in doc.styles:
                            para.style = doc.styles[para_props['style']]
                        else:
                            para.style = para_props['style']
                    except (KeyError, TypeError):
                        pass

                # Alignment (hériter de hAlign de la cellule si pas d'alignment explicite)
                if 'alignment' in para_props:
                    para.alignment = parse_alignment(para_props['alignment'])
                elif 'hAlign' in cell_props:
                    para.alignment = parse_alignment(cell_props['hAlign'])

                # Ajouter les runs
                if 'runs' in para_data:
                    for run_data in para_data['runs']:
                        # Ignorer les page breaks dans les tables
                        if run_data.get('page_break'):
                            continue

                        run_text = run_data.get('text', '')
                        run = para.add_run(run_text)

                        # Appliquer les propriétés du run
                        run_props = run_data.get('properties', {})
                        if run_props.get('bold'):
                            run.bold = True
                        if run_props.get('italic'):
                            run.italic = True

                        # Size (en points)
                        if 'size' in run_props:
                            try:
                                size_half_pt = int(run_props['size'])
                                run.font.size = Pt(size_half_pt / 2)
                            except:
                                pass

                        # Color
                        if 'color' in run_props:
                            try:
                                color_hex = run_props['color']
                                run.font.color.rgb = RGBColor(
                                    int(color_hex[0:2], 16),
                                    int(color_hex[2:4], 16),
                                    int(color_hex[4:6], 16)
                                )
                            except:
                                pass

                        # Font
                        if 'font' in run_props:
                            run.font.name = run_props['font']

                # Texte simple si pas de runs
                elif 'text' in para_data:
                    para.add_run(para_data['text'])

                # Ajouter les paragraphes restants
                for para_data in paragraphs_list[1:]:
                    para = cell.add_paragraph()

                    # Appliquer les propriétés du paragraphe
                    para_props = para_data.get('properties', {})

                    # Style
                    if 'style' in para_props:
                        try:
                            if para_props['style'] in doc.styles:
                                para.style = doc.styles[para_props['style']]
                            else:
                                para.style = para_props['style']
                        except (KeyError, TypeError):
                            pass

                    # Alignment (hériter de hAlign de la cellule si pas d'alignment explicite)
                    if 'alignment' in para_props:
                        para.alignment = parse_alignment(para_props['alignment'])
                    elif 'hAlign' in cell_props:
                        para.alignment = parse_alignment(cell_props['hAlign'])

                    # Ajouter les runs
                    if 'runs' in para_data:
                        for run_data in para_data['runs']:
                            # Ignorer les page breaks dans les tables
                            if run_data.get('page_break'):
                                continue

                            run_text = run_data.get('text', '')
                            run = para.add_run(run_text)

                            # Appliquer les propriétés du run
                            run_props = run_data.get('properties', {})
                            if run_props.get('bold'):
                                run.bold = True
                            if run_props.get('italic'):
                                run.italic = True

                            # Size (en points)
                            if 'size' in run_props:
                                try:
                                    size_half_pt = int(run_props['size'])
                                    run.font.size = Pt(size_half_pt / 2)
                                except:
                                    pass

                            # Color
                            if 'color' in run_props:
                                try:
                                    color_hex = run_props['color']
                                    run.font.color.rgb = RGBColor(
                                        int(color_hex[0:2], 16),
                                        int(color_hex[2:4], 16),
                                        int(color_hex[4:6], 16)
                                    )
                                except:
                                    pass

                            # Font
                            if 'font' in run_props:
                                run.font.name = run_props['font']

                    # Texte simple si pas de runs
                    elif 'text' in para_data:
                        para.add_run(para_data['text'])


def json_to_docx(json_file: str, template_file: str, output_dir: str) -> str:
    """
    Injecte le contenu d'un JSON dans un template DOCX.

    Args:
        json_file (str): Chemin du fichier JSON
        template_file (str): Chemin du template DOCX
        output_dir (str): Répertoire de sortie

    Returns:
        str: Chemin du fichier DOCX créé
    """
    try:
        # Valider l'existence du fichier JSON
        json_path = Path(json_file)
        if not json_path.exists():
            print(f"❌ ERREUR: Le fichier JSON n'existe pas: {json_file}")
            raise FileNotFoundError(f"JSON file not found: {json_file}")

        if not json_path.is_file():
            print(f"❌ ERREUR: La source n'est pas un fichier: {json_file}")
            raise IsADirectoryError(f"JSON source is not a file: {json_file}")

        # Valider l'existence du template
        template_path = Path(template_file)
        if not template_path.exists():
            print(f"❌ ERREUR: Le template n'existe pas: {template_file}")
            raise FileNotFoundError(f"Template file not found: {template_file}")

        print(f"📖 Lecture du fichier JSON: {json_path}")

        # Créer le répertoire de sortie
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        print(f"📁 Répertoire de sortie créé: {output_path}")

        # Charger le JSON
        with open(json_file, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        print(f"✓ JSON chargé avec succès")

        # Charger le template
        print(f"📄 Chargement du template: {template_path}")
        doc = Document(template_file)
        print(f"✓ Template chargé avec succès")

        # Injecter le contenu
        content = json_data.get('document', {}).get('content', [])
        print(f"🔄 Injection de {len(content)} éléments dans le document...")

        for idx, element in enumerate(content, 1):
            elem_type = element.get('type')

            if elem_type == 'Paragraph':
                add_paragraph_from_json(doc, element)

            elif elem_type == 'Table':
                add_table_from_json(doc, element)

        print(f"✓ Tous les éléments ont été injectés")

        # Générer le nom de sortie
        json_stem = Path(json_file).stem.replace('_transformed', '')
        output_file = output_path / f"{json_stem}_generated.docx"

        # Sauvegarder
        print(f"💾 Sauvegarde du fichier DOCX...")
        doc.save(str(output_file))

        # Obtenir la taille du fichier
        file_size = output_file.stat().st_size / 1024  # En KB

        print(f"✅ SUCCÈS: Fichier DOCX généré avec succès!")
        print(f"   📁 Chemin: {output_file}")
        print(f"   💾 Taille: {file_size:.1f} KB")

        return str(output_file)

    except FileNotFoundError as e:
        print(f"❌ ERREUR: {str(e)}")
        raise
    except IsADirectoryError as e:
        print(f"❌ ERREUR: {str(e)}")
        raise
    except Exception as e:
        print(f"❌ ERREUR lors du rendu: {str(e)}")
        raise

def main():
    parser = ArgumentParser(description="Injecte le contenu d'un JSON dans un template DOCX")
    parser.add_argument(
        "-s", "--source_json_file",
        required=True,
        help="Chemin du fichier JSON (transformé)"
    )

    parser.add_argument(
        "-t", "--template",
        default="assets/TEMPLATE.docx",
        help="Chemin du template DOCX (défaut: assets/TEMPLATE.docx)"
    )

    parser.add_argument(
        "-o", "--output_dir",
        default="OUTPUT4_DOCX-RESULT",
        help="Répertoire de sortie (défaut: OUTPUT4_DOCX-RESULT)"
    )

    args = parser.parse_args()

    print("🚀 Démarrage du rendu JSON → DOCX")
    print(f"   📄 Source JSON: {args.source_json_file}")
    print(f"   🎨 Template: {args.template}")
    print(f"   📁 Sortie: {args.output_dir}")
    print()

    try:
        output_file = json_to_docx(args.source_json_file, args.template, args.output_dir)
        print()
        print(f"✨ Rendu terminé avec succès: {output_file}")
    except Exception as e:
        print()
        print(f"❌ Rendu échoué: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()

