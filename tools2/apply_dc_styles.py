#!/usr/bin/env python3
"""
Utilitaire pour appliquer les styles DC à un document.
À intégrer dans parse_reformat.py pour formatage automatique.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path
from typing import Dict, Optional


class DCStyleApplier:
    """Applique les styles DC prédéfinis à un document."""
    
    # Mapping: type de bloc → style DC
    STYLE_MAPPING = {
        'header': 'DC_Header',              # DOSSIER DE COMPÉTENCES
        'name': 'DC_Name',                  # Nom de la personne
        'title': 'DC_Title',                # Data Engineer, Data Analyst
        'section': 'DC_Section',            # Domaines de compétences, Formations, etc.
        'subsection': 'DC_Subsection',      # Gestion de projets, Traitement données, etc.
        'bullet': 'DC_Bullet',              # Points principaux
        'sub_bullet': 'DC_SubBullet',       # Sous-points (ilvl > 0)
        'normal': 'Normal',                 # Texte normal
    }
    
    def __init__(self, template_path: str = './TEMPLATE_DC.docx'):
        """Initialise avec un template contenant les styles."""
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
        
        # Charger le template pour accéder aux styles
        self.template_doc = Document(str(self.template_path))
        print(f"✓ Template chargé: {self.template_path}")
        print(f"✓ Styles disponibles: {len(self.template_doc.styles)}")
    
    def copy_styles_to_document(self, target_doc: Document) -> Document:
        """Copie tous les styles DC du template vers le document cible."""
        try:
            # Copier les styles DC du template
            for source_style in self.template_doc.styles:
                # Chercher les styles commençant par 'DC_'
                if source_style.name.startswith('DC_'):
                    try:
                        # Essayer d'ajouter le style s'il n'existe pas
                        target_doc.styles.add_style(
                            source_style.name,
                            source_style.type
                        )
                    except:
                        # Style existe déjà, c'est ok
                        pass
            
            print(f"✓ Styles DC copiés au document")
            return target_doc
        except Exception as e:
            print(f"✗ Erreur copie styles: {e}")
            return target_doc
    
    def apply_style_to_paragraph(self, paragraph, style_type: str) -> None:
        """Applique un style DC à un paragraphe."""
        style_name = self.STYLE_MAPPING.get(style_type, 'Normal')
        try:
            paragraph.style = style_name
        except:
            print(f"⚠ Style '{style_name}' non trouvé, utilisation de Normal")
            paragraph.style = 'Normal'
    
    def detect_and_apply_styles(self, doc: Document) -> Document:
        """
        Détecte automatiquement le type de paragraphe et applique le style approprié.
        Utilise l'heuristique : texte, taille, alignement.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Heuristique simple
            if not text:
                # Paragraphe vide
                continue
            
            # Chercher le style actuel
            current_style = para.style.name if para.style else 'Normal'
            
            # Appliquer le mappage
            if 'Heading 1' in current_style:
                # Grande section → DC_Section
                para.style = 'DC_Section'
            elif 'Heading 2' in current_style:
                # Sous-section → DC_Subsection
                para.style = 'DC_Subsection'
            elif i < 3 and para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                # Premiers paragraphes centrés = titre
                if 'DOSSIER' in text.upper():
                    para.style = 'DC_Header'
                elif 'DATA' in text.upper() or 'ENGINEER' in text.upper():
                    para.style = 'DC_Title'
                else:
                    # Nom ou autre info centrée
                    para.style = 'DC_Name'
            else:
                # Déterminer le niveau d'indentation (ilvl)
                try:
                    ilvl = para.paragraph_format._element.pPr.numPr.ilvl.val if para.paragraph_format._element.pPr.numPr else 0
                    if ilvl > 0:
                        para.style = 'DC_SubBullet'
                    elif para.paragraph_format.left_indent and para.paragraph_format.left_indent > 0:
                        para.style = 'DC_Bullet'
                    else:
                        para.style = 'Normal'
                except:
                    para.style = 'Normal'
        
        print(f"✓ Styles automatiques appliqués aux {len(doc.paragraphs)} paragraphes")
        return doc


def reformat_with_dc_styles(
    source_doc_path: str,
    output_path: str,
    template_path: str = './TEMPLATE_DC.docx'
) -> None:
    """
    Reformatte un document DC en appliquant les styles prédéfinis.
    
    Usage:
        reformat_with_dc_styles('original.docx', 'reformatted.docx')
    """
    print(f"\n=== REFORMATTAGE AVEC STYLES DC ===\n")
    
    # Charger le document source
    doc = Document(source_doc_path)
    print(f"✓ Document source chargé: {source_doc_path}")
    
    # Initialiser l'applicateur de styles
    applier = DCStyleApplier(template_path)
    
    # Copier les styles DC au document
    doc = applier.copy_styles_to_document(doc)
    
    # Appliquer les styles automatiquement
    doc = applier.detect_and_apply_styles(doc)
    
    # Sauvegarder
    doc.save(output_path)
    print(f"✓ Document reformatté sauvegardé: {output_path}")
    print(f"\n✓ Styles DC appliqués avec succès!")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python apply_dc_styles.py <source.docx> [output.docx]")
        sys.exit(1)
    
    source = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else 'output_with_styles.docx'
    
    reformat_with_dc_styles(source, output)
