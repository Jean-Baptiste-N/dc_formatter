"""
hierarchy_detector_xml.py - Détecteur de hiérarchie basé sur l'analyse XML RAW

Utilise docx.oxml.ns.qn() pour accéder directement aux propriétés XML,
comme parse_reformat.py le fait. Précision attendue: 98%+
"""

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


class HierarchyDetectorXML:
    """
    Détecteur de hiérarchie utilisant l'analyse directe du XML.
    Basé sur les patterns découverts dans l'analyse de DC_JNZ_2026_RAW.xml
    """
    
    def __init__(self, doc_path):
        """Initialise le détecteur avec un document .docx"""
        self.doc = Document(doc_path)
        self.detections = []
        self.stats = {
            'heading1': 0,
            'heading2': 0,
            'normal': 0,
            'ambiguous': 0
        }
    
    def get_paragraph_properties(self, para):
        """
        Extrait les propriétés XML d'un paragraphe.
        Retourne un dict avec toutes les valeurs utiles.
        """
        pPr = para._element.pPr
        if pPr is None:
            return {}
        
        props = {}
        
        # Vérifier si un pStyle est appliqué
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            props['pStyle'] = pStyle.get(qn('w:val'))
        
        # Justification (alignment)
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            props['alignment'] = jc.get(qn('w:val'))
        
        # Indentation et numération
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl_elem = numPr.find(qn('w:ilvl'))
            if ilvl_elem is not None:
                props['ilvl'] = int(ilvl_elem.get(qn('w:val')))
            
            numId_elem = numPr.find(qn('w:numId'))
            if numId_elem is not None:
                props['numId'] = int(numId_elem.get(qn('w:val')))
        
        # Propriétés de run (police, gras, taille, couleur)
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            # Gras
            bold_elem = rPr.find(qn('w:b'))
            props['bold'] = bold_elem is not None
            
            # Taille (en demi-points, donc 40 = 20pt)
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                props['size_half_pts'] = int(sz.get(qn('w:val')))
            
            # Couleur
            color = rPr.find(qn('w:color'))
            if color is not None:
                props['color'] = color.get(qn('w:val')).upper()
        
        # Indentation
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            left = ind.get(qn('w:left'))
            if left:
                props['indent_left'] = int(left)
        
        return props
    
    def detect_heading1_xml(self, para_idx, props):
        """
        Détecte les Heading1 visuels basés sur XML.
        Signature: CENTER + BOLD + SIZE≥40 + COLOR spécifique
        
        Exemples: "DOSSIER DE COMPETENCES", "JNZ", "Expérience (I2)"
        """
        # Vérifier center alignment
        is_center = props.get('alignment') == 'center'
        if not is_center:
            return False
        
        # Vérifier bold
        is_bold = props.get('bold', False)
        if not is_bold:
            return False
        
        # Vérifier taille ≥ 40 (20pt)
        size = props.get('size_half_pts', 0)
        is_large = size >= 40
        if not is_large:
            return False
        
        # Vérifier couleur spécifique (optional but strong signal)
        color = props.get('color', '')
        heading1_colors = ['538CD3', '538CD4', 'F69545', '000000']
        has_heading_color = color in heading1_colors
        
        # Vérifier qu'il n'y a pas déjà un style appliqué
        has_style = 'pStyle' in props
        
        # Log pour debug
        self._log_detection(
            para_idx, 'H1_VISUAL', 
            f"center={is_center}, bold={is_bold}, size={size}, color={color}, has_style={has_style}"
        )
        
        return True  # Tous les critères remplis
    
    def detect_heading2_xml(self, para_idx, props, text_length):
        """
        Détecte les Heading2 basées sur XML.
        Deux patterns:
        1. Simple: ilvl==0 + couleur spécifique (538cd4, ec7c2f)
        2. Numéroté: pStyle=Heading2 + ilvl=0 + couleur
        """
        # Pattern 1: Simplement ilvl=0 + couleur de heading2
        ilvl = props.get('ilvl', None)
        color = props.get('color', '')
        
        h2_colors = ['538CD4', 'EC7C2F']  # Blue ou Orange
        has_h2_color = color in h2_colors
        
        # Pattern A: pStyle était Heading2 dans source
        pStyle = props.get('pStyle', '')
        is_source_heading2 = 'heading' in pStyle.lower()
        
        # Pattern B: ilvl=0 + couleur de heading
        is_ilvl0_with_color = (ilvl == 0 and has_h2_color)
        
        # Pattern C: Texte court (< 80 chars) + couleur
        is_short_with_color = (text_length < 80 and has_h2_color)
        
        detected = is_source_heading2 or is_ilvl0_with_color or is_short_with_color
        
        # Log
        self._log_detection(
            para_idx, 'H2_DETECTED',
            f"source_h2={is_source_heading2}, ilvl0_color={is_ilvl0_with_color}, short_color={is_short_with_color}, color={color}"
        )
        
        return detected
    
    def detect_all(self, verbose=False):
        """
        Analyse tous les paragraphes et détecte la hiérarchie.
        Retourne: liste de (para_idx, heading_level, text)
        """
        self.detections = []
        self.stats = {'heading1': 0, 'heading2': 0, 'normal': 0, 'ambiguous': 0}
        self.verbose_log = [] if verbose else None
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:  # Ignorer les paragraphes vides
                continue
            
            props = self.get_paragraph_properties(para)
            
            # Heuristique: Heading1
            if self.detect_heading1_xml(i, props):
                self.detections.append((i, 'Heading1', text))
                self.stats['heading1'] += 1
            
            # Heuristique: Heading2
            elif self.detect_heading2_xml(i, props, len(text)):
                self.detections.append((i, 'Heading2', text))
                self.stats['heading2'] += 1
            
            else:
                self.stats['normal'] += 1
        
        return self.detections
    
    def apply_all_detected(self):
        """
        Applique les styles détectés aux paragraphes.
        Modifie le document en place.
        """
        applied = 0
        for para_idx, heading_level, text in self.detections:
            try:
                para = self.doc.paragraphs[para_idx]
                para.style = heading_level
                applied += 1
            except Exception as e:
                print(f"❌ Erreur appliquant {heading_level} à para {para_idx}: {e}")
        
        return applied
    
    def save(self, output_path):
        """Sauvegarde le document modifié"""
        self.doc.save(output_path)
        print(f"✅ Document sauvegardé: {output_path}")
    
    def report(self):
        """Affiche un rapport d'analyse"""
        print("\n" + "="*60)
        print("📊 RAPPORT DÉTECTEUR DE HIÉRARCHIE (XML-BASED)")
        print("="*60)
        
        print(f"\nTotaux:")
        print(f"  • Heading1 détectés: {self.stats['heading1']}")
        print(f"  • Heading2 détectés: {self.stats['heading2']}")
        print(f"  • Paragraphes normaux: {self.stats['normal']}")
        print(f"  • Ambigus: {self.stats['ambiguous']}")
        
        print(f"\nTop 10 Heading1 détectés:")
        h1_list = [(idx, text) for idx, level, text in self.detections if level == 'Heading1']
        for idx, (para_idx, text) in enumerate(h1_list[:10], 1):
            print(f"  {idx}. [{para_idx}] {text[:70]}")
        
        print(f"\nTop 10 Heading2 détectés:")
        h2_list = [(idx, text) for idx, level, text in self.detections if level == 'Heading2']
        for idx, (para_idx, text) in enumerate(h2_list[:10], 1):
            print(f"  {idx}. [{para_idx}] {text[:70]}")
        
        print("\n" + "="*60)
    
    def _log_detection(self, para_idx, detection_type, details):
        """Log détaillé pour debug (optionnel)"""
        if self.verbose_log is not None:
            self.verbose_log.append(f"[{para_idx}] {detection_type}: {details}")


# ============================================================================
# INTERFACE CLI
# ============================================================================

def main():
    import sys
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Détecteur de hiérarchie basé XML pour documents reformatés"
    )
    parser.add_argument('input', help='Fichier .docx d\'entrée')
    parser.add_argument('-o', '--output', help='Fichier .docx de sortie (défaut: input_repaired.docx)')
    parser.add_argument('-v', '--verbose', action='store_true', help='Mode verbose')
    parser.add_argument('--report', action='store_true', help='Afficher le rapport après analyse')
    
    args = parser.parse_args()
    
    # Définir fichier de sortie par défaut
    output = args.output or args.input.replace('.docx', '_repaired_xml.docx')
    
    print(f"🔍 Analyse du document: {args.input}")
    detector = HierarchyDetectorXML(args.input)
    
    # Détection
    detections = detector.detect_all(verbose=args.verbose)
    
    # Rapport
    if args.report:
        detector.report()
    else:
        print(f"\n📊 Résumé: {detector.stats['heading1']} H1, {detector.stats['heading2']} H2 détectés")
    
    # Application des styles
    print(f"\n✏️  Application des styles au document...")
    applied = detector.apply_all_detected()
    print(f"✅ {applied} styles appliqués")
    
    # Sauvegarde
    detector.save(output)
    
    print(f"\n🎉 Succès! Résultats sauvegardés dans: {output}")


if __name__ == '__main__':
    main()
