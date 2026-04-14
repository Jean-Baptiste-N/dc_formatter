"""
Détecteur de hiérarchie documentaire pour les CV (DC)
Analyse les signaux visuels et structurels pour assigner des styles Heading1/Heading2
aux paragraphes qui ont perdu leur style lors de la reformatage.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import List, Tuple
import sys


class HierarchyDetector:
    """Détecte et applique la hiérarchie de titres basée sur des heuristiques."""
    
    def __init__(self, doc_path: str):
        self.doc = Document(doc_path)
        self.paragraphs = list(self.doc.paragraphs)
        self.detected = []
    
    def get_font_size_emu(self, run) -> int:
        """Récupère la taille de police en EMU (English Metric Units)."""
        if run.font.size:
            return run.font.size
        return 0
    
    def is_bold(self, paragraph) -> bool:
        """Vérifie si le paragraphe est en bold."""
        for run in paragraph.runs:
            if run.font.bold:
                return True
        return False
    
    def is_centered(self, paragraph) -> bool:
        """Vérifie si le paragraphe est centré."""
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    def get_ilvl(self, paragraph) -> int:
        """Récupère le niveau d'indentation (ilvl)."""
        pPr = paragraph._element.pPr
        if pPr is None:
            return -1
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is not None:
                return int(ilvl.get(qn('w:val')))
        return -1
    
    def get_max_font_size(self, paragraph) -> int:
        """Récupère la taille max de police du paragraphe."""
        max_size = 0
        for run in paragraph.runs:
            if run.font.size and run.font.size > max_size:
                max_size = run.font.size
        return max_size
    
    def detect_heading_level(self, para_idx: int) -> str:
        """
        Détecte le niveau de titre du paragraphe.
        Retourne: "Heading1", "Heading2", ou None
        """
        p = self.paragraphs[para_idx]
        text = p.text.strip()
        
        # Ignorer les paragraphes vides
        if not text or len(text) < 2:
            return None
        
        # Heuristique 1: CENTER + BOLD + GROSSE POLICE = Heading1
        if self.is_centered(p) and self.is_bold(p):
            max_size = self.get_max_font_size(p)
            # 254000 EMU = 20pt, 241300 EMU = 19pt
            if max_size >= 241300:
                return "Heading1"
        
        # Heuristique 2: Texte court centré + bold peut être Heading2
        ilvl = self.get_ilvl(p)
        
        if ilvl == 0 and self.is_bold(p):
            # Texte relativement court (< 80 chars) = probablement un titre
            if len(text) < 80 and not text.startswith(" "):
                # Mais pas si c'est déjà un Heading4
                if p.style.name != "Heading4":
                    return "Heading2"
        
        # Heuristique 3: style="Heading4" → convertir en Heading2
        if p.style.name == "Heading4":
            return "Heading2"
        
        # Heuristique 4: ilvl=0 + texte court + pas de ponctuation ":" à la fin
        # = pourrait être un sous-titre (Heading2)
        if ilvl == 0 and self.is_bold(p) and len(text) < 60:
            if self.has_colon_ending(text):
                return "Heading2"
        
        return None
    
    @staticmethod
    def has_colon_ending(text: str) -> bool:
        """Vérifie si le texte finit par ':' (indicateur de titre/label)."""
        return text.rstrip().endswith(":")
    
    def detect_all(self) -> List[Tuple[int, str, str]]:
        """
        Détecte tous les titres du document.
        Retourne: [(index, niveau, text), ...]
        """
        results = []
        for i, p in enumerate(self.paragraphs):
            level = self.detect_heading_level(i)
            if level:
                results.append((i, level, p.text[:50]))  # premiers 50 chars
                self.detected.append((i, level, p.text))
        return results
    
    def apply_heading_style(self, para_idx: int, level: str) -> None:
        """Applique le style de titre au paragraphe."""
        self.paragraphs[para_idx].style = level
    
    def apply_all_detected(self) -> int:
        """Applique tous les styles détectés. Retourne le nombre appliqué."""
        count = 0
        for para_idx, level, text in self.detected:
            try:
                self.apply_heading_style(para_idx, level)
                count += 1
            except Exception as e:
                print(f"⚠️  Erreur appliquant {level} à '{text}': {e}")
        return count
    
    def save(self, output_path: str) -> None:
        """Sauvegarde le document modifié."""
        self.doc.save(output_path)
    
    def report(self) -> None:
        """Affiche un rapport des détections."""
        print("\n" + "="*70)
        print("RAPPORT DE DÉTECTION DE HIÉRARCHIE")
        print("="*70 + "\n")
        
        heading1_count = 0
        heading2_count = 0
        
        for para_idx, level, text in self.detected:
            if level == "Heading1":
                print(f"  ▼ [Heading1] {text[:60]}")
                heading1_count += 1
            elif level == "Heading2":
                print(f"    ├─ [Heading2] {text[:60]}")
                heading2_count += 1
        
        print(f"\n{'─'*70}")
        print(f"Total:  {heading1_count} × Heading1,  {heading2_count} × Heading2")
        print(f"{'─'*70}\n")
    
    def print_analysis(self, limit: int = 10) -> None:
        """Affiche une analyse détaillée des premiers N paragraphes."""
        print("\n" + "="*70)
        print("ANALYSE HEURISTIQUE (premiers paragraphes)")
        print("="*70 + "\n")
        
        for i in range(min(limit, len(self.paragraphs))):
            p = self.paragraphs[i]
            text = p.text[:40]
            
            props = []
            if self.is_centered(p):
                props.append("CENTER")
            if self.is_bold(p):
                props.append("BOLD")
            ilvl = self.get_ilvl(p)
            if ilvl >= 0:
                props.append(f"ilvl={ilvl}")
            
            max_size = self.get_max_font_size(p)
            if max_size:
                props.append(f"size={max_size}")
            
            detected = self.detect_heading_level(i)
            level_str = f" → {detected}" if detected else ""
            
            print(f"[{i:2d}] {text:45} | {', '.join(props)}{level_str}")
        
        print()


def apply_styles_to_document(input_path: str, output_path: str, verbose: bool = True) -> int:
    """
    Détecte et applique les styles de titre à un document.
    
    Args:
        input_path: Chemin du document source
        output_path: Chemin de sauvegarde
        verbose: Afficher les rapports
    
    Returns:
        Nombre de styles appliqués
    """
    detector = HierarchyDetector(input_path)
    detector.detect_all()
    
    if verbose:
        detector.print_analysis(limit=20)
        detector.report()
    
    count = detector.apply_all_detected()
    detector.save(output_path)
    
    if verbose:
        print(f"✅ Document sauvegardé: {output_path}")
        print(f"✅ Styles appliqués: {count}\n")
    
    return count


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python hierarchy_detector.py <input.docx> [output.docx]")
        sys.exit(1)
    
    input_doc = sys.argv[1]
    output_doc = sys.argv[2] if len(sys.argv) > 2 else input_doc.replace(".docx", "_with_headings.docx")
    
    apply_styles_to_document(input_doc, output_doc)
