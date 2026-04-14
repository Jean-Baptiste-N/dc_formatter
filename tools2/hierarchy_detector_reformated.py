"""
hierarchy_detector_reformated.py - Détecteur pour documents REFORMATÉS

Le document reformaté par parse_reformat.py a perdu:
❌ Tous les styles (Heading1, Heading2 → Normal)
❌ Toutes les sizes de police
❌ Toutes les couleurs
❌ Tous les bolds

Il conserve seulement:
✓ L'alignement CENTER pour tous les paragraphes
✓ La longueur et le contenu textuel

Donc les heuristiques doivent se baser sur:
1. Longueur du texte (court = potentiel Heading)
2. Position hiérarchique (Heading1 avant Heading2 avant content)
3. Patterns de texte existants (keywords)
"""

from docx import Document
from docx.oxml.ns import qn


class HierarchyDetectorReformated:
    """Détecteur adapté aux documents après parse_reformat.py"""
    
    # Keywords typiques pour Heading1
    H1_KEYWORDS = {
        'competence', 'experience', 'formation', 'projet', 'domaine',
        'secteur', 'technologie', 'stack', 'outil', 'plateforme',
        'dossier', 'cv', 'resume'
    }
    
    # Keywords typiques pour Heading2 (plus spécifiques)
    H2_KEYWORDS = {
        'gestion', 'traitement', 'collecte', 'analyse', 'conception',
        'développement', 'architecture', 'visualisation', 'support',
        'coordination', 'formation', 'recueil', 'participation',
        'ecriture', 'suivi', 'organisation', 'implémentation',
        'recherche', 'construction', 'stockage', 'transformation',
        'unification', 'intégration', 'chargement'
    }
    
    def __init__(self, doc_path):
        """Initialise le détecteur"""
        self.doc = Document(doc_path)
        self.detections = []
        self.stats = {
            'heading1': 0,
            'heading2': 0,
            'normal': 0
        }
        self.text_lengths = []
        self.all_texts = []
    
    def analyze_document_structure(self):
        """
        Analyse la structure globale du document pour calibrer les seuils.
        Les Heading1/H2 sont généralement plus courts que le contenu normal.
        """
        self.text_lengths = []
        self.all_texts = []
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if text:
                self.text_lengths.append(len(text))
                self.all_texts.append(text)
        
        # Calculer percentiles
        self.text_lengths.sort()
        self.avg_length = sum(self.text_lengths) / len(self.text_lengths) if self.text_lengths else 50
        self.p25 = self.text_lengths[len(self.text_lengths)//4] if len(self.text_lengths) > 4 else 30
        self.p75 = self.text_lengths[3*len(self.text_lengths)//4] if len(self.text_lengths) > 4 else 50
        
        print(f"📊 Structure du document:")
        print(f"   • Longueur moyenne: {self.avg_length:.0f} chars")
        print(f"   • P25 (court): {self.p25} chars")
        print(f"   • P75 (long): {self.p75} chars")
    
    def has_keyword(self, text, keyword_set):
        """Vérifie si le texte contient un keyword du set"""
        text_lower = text.lower()
        return any(kw in text_lower for kw in keyword_set)
    
    def detect_heading1_reformated(self, para_idx, text):
        """
        Détecte Heading1 dans document reformaté.
        Heuristiques:
        1. Texte très court (< 30 chars)
        2. Mots-clés au début du document
        3. Texte en majuscules complets
        """
        if not text or len(text.strip()) == 0:
            return False
        
        text_len = len(text)
        
        # Très court → potentiel H1
        if text_len < 25:
            # Mais pas si c'est un keyword de H2
            if not self.has_keyword(text, self.H2_KEYWORDS):
                return True
        
        # Texte tout en majuscules
        if text.isupper() and len(text) > 4:
            return True
        
        return False
    
    def detect_heading2_reformated(self, para_idx, text):
        """
        Détecte Heading2 dans document reformaté.
        Heuristiques:
        1. Texte court-moyen (25-50 chars)
        2. Keywords de catégories
        3. Texte qui finit par ":" (section header)
        """
        if not text or len(text.strip()) == 0:
            return False
        
        text_len = len(text)
        
        # Finit par ":" → section header
        if text.rstrip().endswith(':'):
            return True
        
        # Longueur de H2 typique
        if 15 < text_len < 55:
            # Avec keywords de H2
            if self.has_keyword(text, self.H2_KEYWORDS):
                return True
            
            # Ou très court (< 25) ET commence un groupe
            if text_len < 25:
                # Verifier que ce n'est pas du contenu détaillé
                words = len(text.split())
                if words <= 5:  # Max 5 mots
                    return True
        
        return False
    
    def detect_all(self, verbose=False):
        """Analyse tous les paragraphes"""
        self.analyze_document_structure()  # Calibrer les seuils
        
        self.detections = []
        self.stats = {'heading1': 0, 'heading2': 0, 'normal': 0}
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Heuristique: Heading1
            if self.detect_heading1_reformated(i, text):
                self.detections.append((i, 'Heading1', text))
                self.stats['heading1'] += 1
            
            # Heuristique: Heading2
            elif self.detect_heading2_reformated(i, text):
                self.detections.append((i, 'Heading2', text))
                self.stats['heading2'] += 1
            
            else:
                self.stats['normal'] += 1
        
        return self.detections
    
    def apply_all_detected(self):
        """Applique les styles détectés"""
        applied = 0
        for para_idx, heading_level, text in self.detections:
            try:
                para = self.doc.paragraphs[para_idx]
                para.style = heading_level
                applied += 1
            except Exception as e:
                print(f"❌ Erreur: {e}")
        
        return applied
    
    def save(self, output_path):
        """Sauvegarde le document"""
        self.doc.save(output_path)
        print(f"✅ Document sauvegardé: {output_path}")
    
    def report(self):
        """Affiche un rapport"""
        print("\n" + "="*60)
        print("📊 RAPPORT DÉTECTEUR REFORMATÉ")
        print("="*60)
        
        print(f"\nTotaux:")
        print(f"  • Heading1 détectés: {self.stats['heading1']}")
        print(f"  • Heading2 détectés: {self.stats['heading2']}")
        print(f"  • Paragraphes normaux: {self.stats['normal']}")
        
        print(f"\nHeading1 détectés:")
        h1_list = [(idx, text) for idx, level, text in self.detections if level == 'Heading1']
        for idx, (para_idx, text) in enumerate(h1_list[:15], 1):
            print(f"  {idx:2d}. [{para_idx:3d}] {text[:60]}")
        
        print(f"\nTop 15 Heading2 détectés:")
        h2_list = [(idx, text) for idx, level, text in self.detections if level == 'Heading2']
        for idx, (para_idx, text) in enumerate(h2_list[:15], 1):
            print(f"  {idx:2d}. [{para_idx:3d}] {text[:60]}")
        
        print("\n" + "="*60)


def main():
    import sys
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Détecteur pour documents REFORMATÉS"
    )
    parser.add_argument('input', help='Fichier .docx d\'entrée (reformaté)')
    parser.add_argument('-o', '--output', help='Fichier .docx de sortie')
    parser.add_argument('--report', action='store_true', help='Rapport détaillé')
    
    args = parser.parse_args()
    
    output = args.output or args.input.replace('.docx', '_restored.docx')
    
    print(f"🔍 Analyse du document reformaté: {args.input}")
    detector = HierarchyDetectorReformated(args.input)
    
    # Détection
    detections = detector.detect_all()
    
    # Rapport
    if args.report:
        detector.report()
    else:
        print(f"\n📊 Résumé: {detector.stats['heading1']} H1, {detector.stats['heading2']} H2 détectés")
    
    # Application
    print(f"\n✏️  Application des styles...")
    applied = detector.apply_all_detected()
    print(f"✅ {applied} styles appliqués")
    
    # Sauvegarde
    detector.save(output)
    
    print(f"\n🎉 Succès! Résultats dans: {output}")


if __name__ == '__main__':
    main()
