#!/usr/bin/env python3
"""
Crée un template DOCX avec des styles prédéfinis pour le reformattage des Dossiers de Compétences.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_dc_template(output_path: str = "TEMPLATE_DC.docx") -> None:
    """Crée un template DOCX avec styles DC prédéfinis."""
    
    doc = Document()
    
    print("=== CRÉATION DU TEMPLATE DC ===\n")
    
    # =========== STYLES PERSONNALISÉS POUR LE DC ===========
    
    # 1. DC_Header - Titre principal du dossier
    try:
        style = doc.styles.add_style('DC_Header', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(20)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x54, 0x8D, 0xD4)  # Bleu
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("✓ Style 'DC_Header' créé (Titre principal)")
    except Exception as e:
        print(f"✗ Erreur DC_Header: {e}")
    
    # 2. DC_Name - Nom de la personne
    try:
        style = doc.styles.add_style('DC_Name', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(20)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0xEC, 0x7C, 0x30)  # Orange
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("✓ Style 'DC_Name' créé (Nom)")
    except Exception as e:
        print(f"✗ Erreur DC_Name: {e}")
    
    # 3. DC_Title - Titre de poste / fonction
    try:
        style = doc.styles.add_style('DC_Title', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(16)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("✓ Style 'DC_Title' créé (Titre de poste)")
    except Exception as e:
        print(f"✗ Erreur DC_Title: {e}")
    
    # 4. DC_Section - En-tête de grande section (Domaines, Formations, etc.)
    try:
        style = doc.styles.add_style('DC_Section', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(18)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x17, 0x47, 0x78)  # Bleu foncé
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        print("✓ Style 'DC_Section' créé (Section principale)")
    except Exception as e:
        print(f"✗ Erreur DC_Section: {e}")
    
    # 5. DC_Subsection - Sous-section (Gestion de projets, Traitement données, etc.)
    try:
        style = doc.styles.add_style('DC_Subsection', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x30, 0x60, 0x90)  # Bleu moyen
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        print("✓ Style 'DC_Subsection' créé (Sous-section)")
    except Exception as e:
        print(f"✗ Erreur DC_Subsection: {e}")
    
    # 6. DC_Bullet - Points de liste
    try:
        style = doc.styles.add_style('DC_Bullet', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(2)
        print("✓ Style 'DC_Bullet' créé (Point de liste)")
    except Exception as e:
        print(f"✗ Erreur DC_Bullet: {e}")
    
    # 7. DC_SubBullet - Sous-points de liste
    try:
        style = doc.styles.add_style('DC_SubBullet', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.75)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(1)
        print("✓ Style 'DC_SubBullet' créé (Sous-point de liste)")
    except Exception as e:
        print(f"✗ Erreur DC_SubBullet: {e}")
    
    # 8. DC_TableHeader - En-tête de table
    try:
        style = doc.styles.add_style('DC_TableHeader', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # Blanc sur fond bleu
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        print("✓ Style 'DC_TableHeader' créé (En-tête table)")
    except Exception as e:
        print(f"✗ Erreur DC_TableHeader: {e}")
    
    # 9. Modification du style Normal par défaut
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(3)
        print("✓ Style 'Normal' modifié (police: Arial 11pt)")
    except Exception as e:
        print(f"✗ Erreur modification Normal: {e}")
    
    # =========== EXEMPLE D'UTILISATION ===========
    
    print("\n=== AJOUT D'EXEMPLES ===\n")
    
    # Ajouter du contenu d'exemple
    doc.add_paragraph('DOSSIER DE COMPÉTENCES', style='DC_Header')
    doc.add_paragraph('Jean Dupont', style='DC_Name')
    doc.add_paragraph('Data Engineer / Data Analyst', style='DC_Title')
    
    doc.add_paragraph('')  # Espace
    
    doc.add_paragraph('Domaines de Compétences', style='DC_Section')
    doc.add_paragraph('Gestion de Projets', style='DC_Subsection')
    doc.add_paragraph('Recueil des besoins', style='DC_Bullet')
    doc.add_paragraph('Analyse de faisabilité', style='DC_Bullet')
    doc.add_paragraph('Suivi de projets et avancements', style='DC_Bullet')
    
    doc.add_paragraph('Traitement des Données', style='DC_Subsection')
    doc.add_paragraph('Conception de pipelines de données robustes', style='DC_Bullet')
    doc.add_paragraph('Collecte et branchements APIs, ETLs python', style='DC_SubBullet')
    doc.add_paragraph('Chargement sur Data Lakes ou Data Warehouses', style='DC_SubBullet')
    
    # Sauvegarder
    doc.save(output_path)
    print(f"✓ Template sauvegardé : {output_path}")
    print(f"✓ Styles disponibles pour le reformatteur : 9 styles personnalisés")
    print(f"\n=== STYLES CRÉÉS ===")
    print("  1. DC_Header       - Titre principal du DC (20pt, Bleu, Centré)")
    print("  2. DC_Name         - Nom (20pt, Orange, Centré)")
    print("  3. DC_Title        - Titre de poste (16pt, Gras, Centré)")
    print("  4. DC_Section      - Grandes sections (18pt, Bleu foncé)")
    print("  5. DC_Subsection   - Sous-sections (14pt, Bleu moyen, indentée)")
    print("  6. DC_Bullet       - Points de liste (11pt, indentés)")
    print("  7. DC_SubBullet    - Sous-points (10pt, doublement indentés)")
    print("  8. DC_TableHeader  - En-têtes de table (11pt, Gras, Blanc)")
    print("  9. Normal (modifié) - Police standard (Arial 11pt)")


if __name__ == '__main__':
    output_file = './TEMPLATE_DC.docx'
    create_dc_template(output_file)
    print(f"\n✓ Template prêt à utiliser : {output_file}")
